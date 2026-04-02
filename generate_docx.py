#!/usr/bin/env python3
"""
Generate EdGame Living Document (DOCX)
Consolidates all project documentation into a single comprehensive document.
"""

import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────────
BASE = Path(__file__).resolve().parent
DOCS = BASE / "docs"
ADR_FILE = DOCS / "adr" / "001-phase1-architecture.md"
DATA_MODEL_FILE = DOCS / "architecture" / "data-model.surql"
API_ROUTES_FILE = DOCS / "architecture" / "api-routes.md"
MONOREPO_FILE = DOCS / "architecture" / "monorepo-structure.md"
COVERAGE_FILE = DOCS / "assessment" / "coverage-matrix.json"
METRIC_FILE = DOCS / "assessment" / "metric-mapping.json"
ECD_DIR = DOCS / "assessment" / "ecd"
BLUEPRINT_FILE = BASE / "EdGame Analytics Blueprint.md"
DIAGRAMS_DIR = BASE / "TIE204Assignments" / "diagrams"

DIAGRAM_FILES = [
    "system_architecture.html",
    "functional_block_diagram.html",
    "algorithm_mapping.html",
    "algorithm_fsm_combat_resolution.html",
    "algorithm_fsm_adaptive_difficulty.html",
    "algorithm_dataflow_metrics.html",
    "interface_dependency_graph.html",
]

OUTPUT_FILE = BASE / "EdGame_Living_Document.docx"


# ══════════════════════════════════════════════════════════════════════
# STYLES & HELPERS
# ══════════════════════════════════════════════════════════════════════

def setup_styles(doc):
    """Configure document-level styles."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 5):
        sname = f"Heading {level}"
        if sname in doc.styles:
            hs = doc.styles[sname]
            hs.font.name = "Calibri"
            hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
            if level == 1:
                hs.font.size = Pt(22)
                hs.font.bold = True
            elif level == 2:
                hs.font.size = Pt(16)
                hs.font.bold = True
            elif level == 3:
                hs.font.size = Pt(13)
                hs.font.bold = True
            elif level == 4:
                hs.font.size = Pt(11)
                hs.font.bold = True

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def add_page_break(doc):
    doc.add_page_break()


SCREENSHOTS_DIR = BASE / "TIE204Assignments" / "diagrams" / "screenshots"


def add_diagram_placeholder(doc, diagram_name, caption):
    """Insert the diagram screenshot PNG, falling back to a placeholder if missing."""
    png_name = Path(diagram_name).stem + ".png"
    png_path = SCREENSHOTS_DIR / png_name

    if png_path.exists():
        # Insert the actual image, scaled to page width
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(png_path), width=Inches(6.2))
    else:
        # Fallback: gray placeholder box
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        cell.width = Inches(5.5)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n\n[INSERT SCREENSHOT: {diagram_name}]\n\n")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.italic = True

    # Caption below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure: {caption}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()  # spacer


def _apply_inline_formatting(paragraph, text):
    """Parse inline markdown formatting (bold, italic, inline code) and add runs."""
    # Pattern matches: **bold**, *italic*, `code`, or plain text
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))')
    for match in pattern.finditer(text):
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif match.group(5):  # plain text
            paragraph.add_run(match.group(5))


def add_markdown_content(doc, text, heading_offset=0):
    """
    Parse markdown text and convert to python-docx elements.
    Handles: headings, bullets, numbered lists, tables, code blocks,
    blockquotes, horizontal rules, and inline formatting.
    heading_offset shifts heading levels (e.g., offset=1 makes # → Heading 2).
    """
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Code block ──
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            # Add code block as a styled paragraph
            code_text = "\n".join(code_lines)
            if code_text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # Add background shading
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                p._p.get_or_add_pPr().append(shading)
            continue

        # ── Table ──
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)
            continue

        # ── Heading ──
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1)) + heading_offset
            level = max(1, min(level, 4))  # clamp 1-4
            title = heading_match.group(2).strip()
            # Remove markdown links from headings
            title = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', title)
            doc.add_heading(title, level=level)
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r'^(\*{3,}|-{3,}|_{3,})\s*$', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin line
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # ── Blockquote ──
        if line.strip().startswith(">"):
            quote_text = line.strip().lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── Bullet list ──
        bullet_match = re.match(r'^(\s*)[*\-+]\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            content = bullet_match.group(2).strip()
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 4:
                p.paragraph_format.left_indent = Cm(2)
            elif indent >= 2:
                p.paragraph_format.left_indent = Cm(1.5)
            p.clear()
            _apply_inline_formatting(p, content)
            i += 1
            continue

        # ── Numbered list ──
        num_match = re.match(r'^(\s*)\d+[.)]\s+(.*)', line)
        if num_match:
            content = num_match.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            p.clear()
            _apply_inline_formatting(p, content)
            i += 1
            continue

        # ── Footnote reference lines (skip) ──
        if re.match(r'^\[\^\d+\]:', line):
            i += 1
            continue

        # ── Empty line ──
        if not line.strip():
            i += 1
            continue

        # ── Regular paragraph ──
        p = doc.add_paragraph()
        _apply_inline_formatting(p, line.strip())
        i += 1


def _add_markdown_table(doc, table_lines):
    """Convert markdown table lines to a docx table."""
    # Parse rows, skip separator row
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip separator rows (---, :--:, etc.)
        if all(re.match(r'^[\-:]+$', c) for c in cells if c):
            continue
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    # Pad rows to have equal columns
    for r in rows:
        while len(r) < num_cols:
            r.append("")

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            _apply_inline_formatting(p, cell_text)
            if ri == 0:
                # Header row styling
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                for run in p.runs:
                    run.font.size = Pt(10)
                if ri % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FA"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # spacer


def read_file(path):
    """Read file content, return empty string if not found."""
    try:
        return Path(path).read_text(encoding="utf-8")
    except Exception as e:
        return f"[Error reading {path}: {e}]"


# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════

def write_title_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EdGame Analytics Platform")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Technical Architecture & Design Document")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Living Document — Version 1.0")
    run.font.size = Pt(13)
    run.font.italic = True

    doc.add_paragraph()

    for info in [
        "Author: Yousef Radwan",
        "Course: TIE 251 — Capstone Computing Studies",
        "Institution: KAUST",
        "Date: March 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(info)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════

def write_executive_summary(doc):
    doc.add_heading("1. Executive Summary", level=1)

    content = """\
EdGame is a game-based learning analytics platform designed for K-12 classrooms. It integrates five educational games with a real-time analytics dashboard, enabling teachers to gain actionable insights into student knowledge, engagement, strategic thinking, collaboration, and social-emotional development.

The platform is built on **Evidence-Centered Design (ECD)**, using stealth assessment and **Bayesian Knowledge Tracing (BKT)** to measure student competencies without interrupting gameplay. Every game mechanic is intentionally designed to produce observable evidence mapped to specific knowledge components and analytics dimensions.

**Key capabilities:**
- **5 educational games** spanning math and science, each targeting different analytics dimensions
- **50+ learning metrics** across 6 analytical dimensions (Cognitive, Behavioral, Strategic, Social, Affective, Temporal)
- **Real-time teacher dashboards** with class-level and student-level views
- **Adaptive difficulty** that adjusts question difficulty based on estimated mastery
- **Privacy-first architecture** compliant with FERPA and COPPA regulations

**Technology stack:** Next.js 14, SurrealDB, KAPLAY.js, pnpm + Turborepo monorepo, deployed on Railway.

This document serves as the canonical technical reference for the EdGame project, consolidating architecture decisions, data models, API specifications, assessment frameworks, and game design documentation into a single living document.
"""
    add_markdown_content(doc, content)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 2: SYSTEM OVERVIEW & VISION
# ══════════════════════════════════════════════════════════════════════

def write_system_overview(doc):
    doc.add_heading("2. System Overview & Vision", level=1)

    content = """\
## Product Description

EdGame transforms classroom assessment by embedding rigorous measurement within engaging game experiences. Rather than testing students with traditional quizzes, EdGame captures rich behavioral telemetry during gameplay and uses evidence-centered design to infer mastery of specific knowledge components.

## Strategic Principle

**"Every game mechanic is an assessment opportunity."** No mechanic exists purely for entertainment — each interaction generates observable evidence that feeds into the analytics pipeline.

## Design Principles

- **Stealth Assessment:** Students are assessed through natural gameplay, not separate test modes
- **Evidence-Centered Design:** Every observable maps to specific knowledge components through explicit task and evidence models
- **Teacher Empowerment:** Analytics dashboards provide actionable insights, not just raw data
- **Adaptive Learning:** Difficulty adjusts in real-time based on Bayesian mastery estimates
- **Privacy by Design:** Minimal data collection, FERPA/COPPA compliance, no PII in telemetry

## 5-Game Portfolio

| Game | Genre | Subject | Primary Dimension | Status |
|------|-------|---------|-------------------|--------|
| **Pulse Realms** | 3v3 Team Arena | Math & Science | D4: Social & Collaborative | Active Development |
| **Concept Cascade** | Tower Defense | Mathematics | D3: Strategic Behavior | Planned |
| **Lab Explorer** | Virtual Science Lab | Science | D3: Strategic Behavior | Planned |
| **Knowledge Quest** | Turn-Based RPG | Math & Science | D5: Affective & SEL | Planned |
| **Survival Equation** | Cooperative Puzzle | Math & Science | D4: Social & Collaborative | Planned |
"""
    add_markdown_content(doc, content)

    add_diagram_placeholder(doc, "system_architecture.html",
                            "System Architecture — High-level overview of EdGame platform components")
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 3: ARCHITECTURE DECISIONS (ADR)
# ══════════════════════════════════════════════════════════════════════

def write_adr_section(doc):
    doc.add_heading("3. Architecture Decisions (ADR)", level=1)

    p = doc.add_paragraph()
    _apply_inline_formatting(p, "This section contains the full Architecture Decision Record (ADR) for Phase 1, documenting all binding technology decisions with their rationale, consequences, and considered alternatives.")
    doc.add_paragraph()

    content = read_file(ADR_FILE)
    # Skip the top-level heading since we already added one
    content = re.sub(r'^#\s+.*\n', '', content, count=1)
    add_markdown_content(doc, content, heading_offset=1)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 4: SYSTEM ARCHITECTURE & ALGORITHMS
# ══════════════════════════════════════════════════════════════════════

def write_architecture_algorithms(doc):
    doc.add_heading("4. System Architecture & Algorithms", level=1)

    intro = """\
This section details the functional architecture, algorithm designs, and interface specifications for the EdGame platform. Each subsection includes a diagram placeholder for the corresponding visual from the TIE 204 assignments.

## Functional Block Diagram
"""
    add_markdown_content(doc, intro)
    add_diagram_placeholder(doc, "functional_block_diagram.html",
                            "Functional Block Diagram — Major system components and data flow")

    doc.add_heading("Algorithm Mapping", level=2)
    p = doc.add_paragraph()
    _apply_inline_formatting(p, "Maps game mechanics to assessment algorithms and shows how player actions translate into analytics metrics through the ECD pipeline.")
    add_diagram_placeholder(doc, "algorithm_mapping.html",
                            "Algorithm Mapping — Game mechanics to assessment algorithms")

    doc.add_heading("Combat Resolution FSM", level=2)
    p = doc.add_paragraph()
    _apply_inline_formatting(p, "The Combat Resolution Finite State Machine governs how player actions (attack, heal, shield) are resolved in Pulse Realms. Each state transition emits telemetry events that feed into the analytics pipeline.")
    add_diagram_placeholder(doc, "algorithm_fsm_combat_resolution.html",
                            "Combat Resolution FSM — State transitions during Pulse Realms combat")

    doc.add_heading("Adaptive Difficulty FSM", level=2)
    content = """\
The Adaptive Difficulty system uses Bayesian Knowledge Tracing to estimate student mastery of specific knowledge components. The FSM manages transitions between difficulty levels based on mastery estimates:

- **Easy Zone (mastery < 0.3):** Foundation-level questions to build confidence
- **Growth Zone (0.3 ≤ mastery < 0.7):** Grade-level questions for active learning
- **Challenge Zone (mastery ≥ 0.7):** Above-grade questions to push growth
- **Transitions** require sustained performance (not single answers) to prevent oscillation
"""
    add_markdown_content(doc, content)
    add_diagram_placeholder(doc, "algorithm_fsm_adaptive_difficulty.html",
                            "Adaptive Difficulty FSM — BKT-driven difficulty transitions")

    doc.add_heading("Metric Computation Dataflow", level=2)
    p = doc.add_paragraph()
    _apply_inline_formatting(p, "Shows the dataflow from raw telemetry events through aggregation and computation to final metric values displayed on teacher dashboards. Covers all 6 analytics dimensions.")
    add_diagram_placeholder(doc, "algorithm_dataflow_metrics.html",
                            "Metric Computation Dataflow — From raw events to dashboard metrics")

    doc.add_heading("Interface Dependency Graph", level=2)
    p = doc.add_paragraph()
    _apply_inline_formatting(p, "Documents the TypeScript interfaces and their dependencies across the shared packages, game clients, and API layer. Ensures type safety across the monorepo.")
    add_diagram_placeholder(doc, "interface_dependency_graph.html",
                            "Interface Dependency Graph — TypeScript interfaces across packages")

    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 5: SPACETIMEDB FOR MULTIPLAYER (NEW)
# ══════════════════════════════════════════════════════════════════════

def write_spacetimedb_section(doc):
    doc.add_heading("5. SpacetimeDB for Multiplayer", level=1)

    content = """\
## What is SpacetimeDB?

SpacetimeDB is a server-side database engine that combines a relational database with a real-time multiplayer server. Instead of writing separate game server code and database queries, developers write **server-side modules** (in Rust or C#) that define tables, reducers (server functions), and subscriptions. Clients connect via WebSocket and subscribe to table rows — when data changes, all subscribed clients receive automatic delta updates with no manual networking code required.

**Key architectural properties:**
- **Single source of truth:** Game state lives in the database; there is no separate "game server" process
- **Automatic replication:** Clients declare which rows they care about; SpacetimeDB pushes changes
- **Conflict resolution built-in:** Reducers execute transactionally on the server, eliminating race conditions
- **Sub-50ms latency:** Optimized for real-time multiplayer games
- **Zero server code for networking:** No WebSocket handlers, no serialization, no message routing

## Why SpacetimeDB Fits EdGame

EdGame's multiplayer games (Pulse Realms 3v3, Survival Equation 2-4 player co-op) require **real-time synchronization** of game state, **server-authoritative logic** (to prevent cheating and ensure valid telemetry), and **persistent state** (match history, player stats). SpacetimeDB unifies all three requirements in a single system.

For an educational platform, the key advantage is **telemetry integrity**: because all game state mutations happen server-side in reducers, every action is automatically logged with a server timestamp. This eliminates the risk of client-side telemetry tampering and ensures assessment data is trustworthy.

## Comparison with Alternatives

| Feature | SpacetimeDB | Photon (PUN/Fusion) | Colyseus | Custom WebSocket | Firebase Realtime |
|---------|------------|---------------------|----------|-----------------|-------------------|
| **Architecture** | DB + Server in one | Cloud relay/host | Node.js rooms | Manual server | Cloud NoSQL |
| **Real-time sync** | Automatic (subscriptions) | Built-in (RPCs/state) | Built-in (schema sync) | Manual | Automatic (listeners) |
| **Server authority** | Yes (reducers) | Optional (host mode) | Yes (room handlers) | Yes (custom) | No (client writes) |
| **Persistence** | Built-in (is a database) | External DB needed | External DB needed | External DB needed | Built-in |
| **Language** | Rust/C# modules | C#/Unity | TypeScript/Node.js | Any | JavaScript SDK |
| **Latency** | <50ms | <100ms (relay) | <100ms | Depends | 100-500ms |
| **Cost model** | Self-host or cloud | Per-CCU pricing | Self-host | Self-host | Pay per read/write |
| **Telemetry integrity** | Inherent (server-only mutations) | Manual implementation | Manual implementation | Manual implementation | Weak (client can write) |
| **Learning curve** | Moderate (new paradigm) | Low (Unity-native) | Low (Node.js) | High | Low |
| **Offline support** | Limited | Photon has reconnect | Manual | Manual | Strong |

## Pros and Cons for EdGame

### Pros

- **Unified data layer:** No impedance mismatch between game state and persistent storage — SpacetimeDB is both
- **Telemetry integrity by design:** All mutations are server-side reducers, guaranteeing authentic assessment data
- **Automatic real-time sync:** Drastically reduces multiplayer networking code complexity
- **Transactional game logic:** Reducers execute atomically, preventing race conditions in combat resolution and scoring
- **Natural fit for ECD pipeline:** Server-side events can be directly fed to the analytics pipeline without a separate collection layer
- **Self-hostable:** Important for school deployments with data sovereignty requirements

### Cons

- **Newer technology:** Smaller community, fewer production references compared to Photon/Colyseus
- **Rust module development:** Requires Rust proficiency for server modules (mitigated by C# alternative)
- **Not a full analytics DB:** Still need SurrealDB for long-term analytics aggregation, dashboard queries, and complex graph traversals
- **Limited ecosystem:** No built-in matchmaking, lobby system, or voice chat — must be implemented manually
- **Vendor dependency risk:** Newer project with less established long-term support guarantees
- **Browser client maturity:** TypeScript/WASM client SDK is newer than native SDKs

## Integration Architecture: SpacetimeDB + SurrealDB

EdGame uses a **dual-database architecture**:

- **SpacetimeDB** handles real-time game state during active matches (player positions, health, combat state, active questions, team composition). This is the "hot" layer — fast, transactional, and ephemeral per match.

- **SurrealDB** handles persistent analytics data, user accounts, curriculum mapping, dashboard queries, and long-term metric aggregation. This is the "warm/cold" layer — optimized for complex queries and graph relationships.

**Data flow:**
1. Match starts → SpacetimeDB creates match tables, players subscribe
2. During gameplay → All actions execute as SpacetimeDB reducers, state syncs to all clients
3. Reducers emit telemetry events → A bridge service streams events to SurrealDB in near-real-time
4. Match ends → Final match summary written to SurrealDB; SpacetimeDB match data archived or purged
5. Dashboard queries → Read exclusively from SurrealDB

## Games Requiring SpacetimeDB

| Game | Players | SpacetimeDB Role | Priority |
|------|---------|------------------|----------|
| **Pulse Realms** | 3v3 (6 players) | Full real-time combat sync, team coordination | Phase 1 |
| **Survival Equation** | 2-4 co-op | Shared puzzle state, role-specific information hiding | Phase 2 |
| **Knowledge Quest** | 1 (single-player) | Not needed — client-side with REST API | N/A |
| **Concept Cascade** | 1 (single-player) | Not needed — client-side with REST API | N/A |
| **Lab Explorer** | 1 (potential co-op later) | Future consideration for collaborative experiments | Phase 3 |

## Phased Rollout Plan

**Phase 1 (Current):** Pulse Realms uses a simplified client-authoritative model with REST API telemetry submission. This works for classroom LAN environments where cheating risk is low.

**Phase 2:** Integrate SpacetimeDB for Pulse Realms multiplayer. Migrate combat resolution, team state, and question delivery to server-side reducers. Implement the SpacetimeDB → SurrealDB bridge service.

**Phase 3:** Build Survival Equation on SpacetimeDB from the start, using the architecture established in Phase 2. Implement information hiding (role-specific data only visible to each player's subscription).

**Phase 4:** Evaluate SpacetimeDB for collaborative Lab Explorer experiments if co-op features are added.
"""
    add_markdown_content(doc, content)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 6: AI-POWERED QUESTION GENERATION PIPELINE (NEW)
# ══════════════════════════════════════════════════════════════════════

def write_question_generation_section(doc):
    doc.add_heading("6. AI-Powered Question Generation Pipeline", level=1)

    content = """\
## Overview

A critical requirement for EdGame is generating high-quality, curriculum-aligned questions at scale. Teachers should not need to manually author hundreds of MCQs — instead, they upload their curriculum materials (textbook chapters, lesson plans, worksheets) and the platform generates questions automatically, mapped to the appropriate knowledge components.

This section describes the end-to-end pipeline from teacher upload to question delivery within games.

## Pipeline Architecture

The question generation pipeline consists of six stages:

### Stage 1: Upload & Ingest

Teachers upload curriculum materials through the dashboard. Supported formats include PDF, DOCX, PPTX, and plain text. The system extracts text content using format-specific parsers, preserving structural information (headings, sections, figures with captions).

**Key considerations:**
- OCR fallback for scanned PDFs
- Chunking strategy: split by section headings, with overlap to preserve context
- Metadata extraction: grade level, subject, chapter/unit identification
- Storage: raw files in object storage, extracted text + metadata in SurrealDB

### Stage 2: Content Analysis & Knowledge Component Extraction

Before generating questions, the system analyzes the uploaded content to identify discrete knowledge components (KCs) that can be assessed. This step uses an LLM with a structured output schema.

**Prompt strategy:**
- System prompt defines the KC extraction task with examples
- Input: chunked text content + grade level + subject area
- Output: structured JSON with KC ID, description, Bloom's taxonomy level, prerequisite KCs
- Post-processing: deduplicate KCs, merge near-duplicates, link to existing KC taxonomy

### Stage 3: Question Generation

For each identified KC, the system generates multiple-choice questions using an LLM with carefully designed prompts.

**Algorithm details:**

- **Prompt template** includes: the source text chunk, the target KC, the desired difficulty level (1-5), the desired Bloom's taxonomy level, and 2-3 example questions as few-shot demonstrations
- **Difficulty calibration:** Difficulty is parameterized along multiple axes:
  - *Conceptual complexity:* recall vs. application vs. analysis
  - *Distractor similarity:* how close wrong answers are to the correct answer
  - *Context dependency:* whether the question requires integrating multiple concepts
  - *Linguistic complexity:* vocabulary level and sentence structure
- **Output format:** Structured JSON with question stem, 4 options (1 correct, 3 distractors), correct answer index, KC mapping, difficulty rating, Bloom's level, and source reference

**Distractor generation strategy:**
- **Common misconception distractors:** Based on known student misconceptions for the topic (sourced from educational research databases)
- **Computational error distractors:** For math questions, include answers that result from common procedural errors (e.g., sign errors, order-of-operations mistakes)
- **Partial knowledge distractors:** Answers that would be correct if the student only partially understood the concept
- **Plausible but unrelated distractors:** Factually correct statements that don't answer the specific question

### Stage 4: Quality Filtering

Generated questions pass through automated quality checks before entering the question bank:

- **Answer validation:** Verify the marked correct answer is actually correct (cross-check with a second LLM call)
- **Distractor quality:** Ensure distractors are plausible but unambiguously wrong
- **Readability check:** Flesch-Kincaid grade level appropriate for target audience
- **Bias detection:** Flag questions with cultural, gender, or socioeconomic bias indicators
- **Duplicate detection:** Semantic similarity check against existing question bank
- **Bloom's alignment:** Verify the question actually assesses the claimed Bloom's level

Questions that fail any check are flagged for human review or regenerated with adjusted prompts.

### Stage 5: KC Mapping & Difficulty Tagging

Validated questions are mapped to the EdGame knowledge component taxonomy and tagged with metadata for the adaptive difficulty system:

- **KC mapping:** Link to existing KCs in the taxonomy, or create new KCs if the curriculum introduces novel concepts
- **Difficulty parameters:** Initial difficulty estimate (refined through BKT as students answer)
- **Prerequisite chain:** Which KCs must be mastered before this question is appropriate
- **Game integration tags:** Which games can use this question (based on format, timing, and subject)

### Stage 6: Storage & Delivery

Questions are stored in SurrealDB with full metadata and served to games through the question engine API:

- **Question bank structure:** Organized by subject → grade → unit → KC → difficulty
- **Versioning:** Questions are versioned; retired questions are soft-deleted to preserve historical analytics
- **Delivery API:** Games request questions by KC + difficulty range; the engine selects based on BKT mastery estimates and spacing algorithms

## Nuances & Improvements

### Bloom's Taxonomy Integration

Questions are tagged with Bloom's taxonomy levels (Remember, Understand, Apply, Analyze, Evaluate, Create). The adaptive difficulty system can target specific Bloom's levels based on student mastery:

- **Low mastery (< 0.3):** Primarily Remember and Understand questions
- **Medium mastery (0.3-0.7):** Apply and Analyze questions
- **High mastery (> 0.7):** Evaluate and Create questions (where applicable for MCQ format)

### Domain-Specific Generation

Different subjects require different generation strategies:

- **Mathematics:** Questions must include precise notation, step-by-step solution paths, and computation-based distractors. The LLM is prompted with LaTeX formatting guidelines.
- **Science:** Questions should reference experimental contexts, include diagram descriptions, and test conceptual understanding rather than memorization.
- **Mixed (game-contextualized):** Questions embedded in game narratives must feel natural within the game world while maintaining educational rigor.

### Teacher Review Workflow

While the pipeline is automated, teachers maintain control:

1. **Preview generated questions** before they enter the active question bank
2. **Edit questions** directly (corrections auto-propagate to metadata)
3. **Rate question quality** (feedback improves future generation)
4. **Add custom questions** that bypass the generation pipeline
5. **Set curriculum scope** to control which topics generate questions

### Versioning & Analytics Feedback Loop

- Questions that consistently show high discrimination (good at separating high/low mastery students) are rated higher and served more frequently
- Questions with poor discrimination are flagged for review or retirement
- BKT parameters for each question are continuously refined as more students answer them
- Generation prompts are iteratively improved based on quality metrics of generated questions

### Adaptive Tagging

As students interact with questions in games, the system refines its understanding of each question's properties:

- **Empirical difficulty:** Actual p-value (% correct) compared to LLM-estimated difficulty
- **Discrimination index:** How well the question differentiates between high and low mastery students
- **Time characteristics:** Average response time, time-pressure sensitivity
- **Distractor analysis:** Which distractors are selected most often, indicating specific misconceptions

## Integration with Game Systems

The question generation pipeline connects to two game systems:

**Question Engine:** Receives requests from games for questions matching specific criteria (KC, difficulty, Bloom's level) and returns appropriately selected questions using spacing and mastery-aware algorithms.

**Adaptive Difficulty System:** Uses BKT mastery estimates to determine the appropriate difficulty level for the next question. As mastery increases, the system requests harder questions with higher Bloom's levels. If mastery drops, it reverts to reinforcement-level questions.
"""
    add_markdown_content(doc, content)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 7: TABLET VS PC EXPERIMENTATION (NEW)
# ══════════════════════════════════════════════════════════════════════

def write_tablet_vs_pc_section(doc):
    doc.add_heading("7. Tablet vs PC Experimentation", level=1)

    content = """\
## Rationale

Saudi Arabian K-12 classrooms vary significantly in available technology. Some schools have dedicated computer labs with desktop PCs, while others rely on tablet carts (primarily iPads and Android tablets). EdGame must understand how **form factor affects learning outcomes and engagement** to make informed deployment recommendations and ensure equitable assessment across devices.

This section outlines the experimental design for comparing tablet and PC gameplay experiences using a single EdGame title.

## Game Selection for Cross-Platform Testing

**Pulse Realms** is the primary candidate for tablet adaptation because:

- It is the most developed game (active development, Phase 1 priority)
- Its touch-based combat mechanics (tap to select ability, tap target) translate naturally to touchscreens
- The 3v3 format creates consistent, controlled match conditions for comparison
- It generates the richest telemetry data across multiple analytics dimensions

**Adaptation scope:** The game already runs in a web browser (KAPLAY.js); the tablet version requires responsive UI scaling and touch input optimization, not a full rewrite.

## Key Differences: Tablet vs PC

| Dimension | PC (Desktop/Laptop) | Tablet (iPad/Android) |
|-----------|---------------------|----------------------|
| **Input method** | Mouse + keyboard (precise, fast) | Touch (direct, less precise) |
| **Screen size** | 13-24 inches | 9-11 inches |
| **UI density** | Can display more information | Requires larger touch targets |
| **Typing** | Physical keyboard (fast text input) | On-screen keyboard (slower, covers screen) |
| **Posture** | Seated at desk, fixed screen | Handheld or propped, variable angle |
| **Distraction** | Browser tabs, other apps | Notifications, app switching |
| **Accessibility** | Mouse accommodation, screen readers | VoiceOver/TalkBack, Switch Access |
| **Network** | Typically wired/stable Wi-Fi | Wi-Fi only, potentially less stable |

## Adaptation Requirements

### UI/UX Modifications for Tablet

- **Touch targets:** Minimum 44x44 pt for all interactive elements (Apple HIG standard)
- **Responsive layout:** Game UI scales to tablet viewport; HUD elements repositioned to avoid thumb zones
- **Gesture support:** Swipe to switch abilities, pinch to zoom game map (optional)
- **Input fallback:** MCQ answers are tap-only (no keyboard input required)
- **Orientation lock:** Landscape mode enforced during gameplay
- **Performance:** Target 30fps minimum on iPad 7th gen / Samsung Galaxy Tab A8

### Telemetry Additions

Additional telemetry events for the tablet version:

- `input_method`: "touch" vs "mouse" tag on all action events
- `touch_precision`: distance between touch point and target center (px)
- `response_area_size`: actual rendered size of answer buttons (px)
- `orientation_changes`: count of orientation switches during match
- `device_model`: specific device identification for performance correlation

## Study Design

### Participants

- **Target:** 60-120 students, grades 6-8, from 2-4 Saudi Arabian schools
- **Assignment:** Within-subject crossover design — each student plays on both PC and tablet (counterbalanced order to control for learning effects)
- **Duration:** 2 matches per device, spaced 1 week apart

### Experimental Conditions

| Condition | Device | Input | Setting |
|-----------|--------|-------|---------|
| **PC** | School desktop or laptop | Mouse + keyboard | Computer lab |
| **Tablet** | iPad or Android tablet | Touchscreen | Classroom (tablet cart) |

### Outcome Measures

**Engagement metrics (from telemetry):**
- Actions per minute (interaction rate)
- Session duration and voluntary continuation
- Off-task behavior indicators (inactivity periods)
- Emotional engagement proxies (ability selection speed, celebration actions)

**Performance metrics (from analytics pipeline):**
- Question accuracy (overall and by KC)
- Response time distribution
- BKT mastery progression rate
- Strategic depth (ability usage patterns, team coordination)

**Usability metrics (from survey + telemetry):**
- System Usability Scale (SUS) adapted for age group
- Touch precision data (tablet only)
- Self-reported preferences and comfort
- Error rates (misclicks, unintended actions)

### Analysis Plan

- **Primary analysis:** Paired t-tests / Wilcoxon signed-rank tests comparing engagement and performance metrics across conditions
- **Secondary analysis:** Interaction effects with student characteristics (prior gaming experience, device familiarity, gender)
- **Equivalence testing:** TOST procedure to determine if tablet assessment is equivalent to PC assessment (important for fairness)
- **Qualitative:** Thematic analysis of open-ended survey responses

## Cross-Platform Assessment Fairness

A critical question is whether **assessment scores are comparable across devices**. If tablet users systematically score lower due to input difficulty rather than knowledge gaps, the platform must account for this.

**Fairness safeguards:**

- **Device-adjusted timing:** Response time thresholds adjusted for input method (touch is inherently slower than mouse click)
- **Difficulty calibration per device:** If significant performance differences exist, BKT parameters can be calibrated separately for each input method
- **DIF analysis:** Differential Item Functioning analysis to identify questions that are unfairly harder on one device
- **Minimum hardware requirements:** Establish and communicate minimum device specifications to prevent performance-related score depression
- **Teacher visibility:** Dashboard shows device type per session so teachers can contextualize results

## Expected Outcomes

This experiment will produce:
1. **Deployment recommendations:** Data-driven guidance on which devices are suitable for EdGame
2. **Device-specific UX guidelines:** Optimizations for tablet gameplay based on empirical data
3. **Assessment equivalence evidence:** Statistical evidence of whether tablet and PC assessments are comparable
4. **Fairness adjustments:** Calibration parameters if device-specific scoring adjustments are needed
5. **Publication-ready findings:** Contribution to the educational technology literature on form factor effects in game-based assessment
"""
    add_markdown_content(doc, content)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 8: DATA MODEL
# ══════════════════════════════════════════════════════════════════════

def write_data_model(doc):
    doc.add_heading("8. Data Model", level=1)

    p = doc.add_paragraph()
    _apply_inline_formatting(p, "The complete SurrealDB data model for EdGame Phase 1. Defines all tables, fields, permissions, and server-side functions.")
    doc.add_paragraph()

    content = read_file(DATA_MODEL_FILE)
    # Render as a code block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(content)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._p.get_or_add_pPr().append(shading)

    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 9: REST API SPECIFICATION
# ══════════════════════════════════════════════════════════════════════

def write_api_spec(doc):
    doc.add_heading("9. REST API Specification", level=1)

    p = doc.add_paragraph()
    _apply_inline_formatting(p, "Complete Phase 1 API routes for the EdGame platform. All endpoints are Next.js Route Handlers deployed on Railway.")
    doc.add_paragraph()

    content = read_file(API_ROUTES_FILE)
    content = re.sub(r'^#\s+.*\n', '', content, count=1)
    add_markdown_content(doc, content, heading_offset=1)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 10: MONOREPO STRUCTURE
# ══════════════════════════════════════════════════════════════════════

def write_monorepo_structure(doc):
    doc.add_heading("10. Monorepo Structure", level=1)

    p = doc.add_paragraph()
    _apply_inline_formatting(p, "pnpm workspaces + Turborepo monorepo structure. See ADR-001 (Section 3) for rationale.")
    doc.add_paragraph()

    content = read_file(MONOREPO_FILE)
    content = re.sub(r'^#\s+.*\n', '', content, count=1)
    add_markdown_content(doc, content, heading_offset=1)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 11: ASSESSMENT FRAMEWORK — ECD
# ══════════════════════════════════════════════════════════════════════

def write_assessment_framework(doc):
    doc.add_heading("11. Assessment Framework — Evidence-Centered Design", level=1)

    intro = """\
## ECD Overview

Evidence-Centered Design (ECD) is the foundational assessment architecture for EdGame. It provides a principled framework for designing assessments that are embedded within gameplay, ensuring that every game mechanic produces interpretable evidence about student competencies.

ECD consists of four interconnected models:

- **Student Model:** What knowledge, skills, and attributes are we measuring? Defined as Knowledge Components (KCs) with mastery probabilities estimated via Bayesian Knowledge Tracing (BKT).
- **Task Model:** What game situations elicit evidence? Defined as specific game mechanics, scenarios, and interactions that create observable opportunities.
- **Evidence Model:** How do we interpret observations? Defined as rules that connect observable behaviors (e.g., answer correctness, response time, strategy choices) to updates in the student model.
- **Assembly Model:** How do we select and sequence tasks? Managed by the adaptive difficulty system and question engine.

### Stealth Assessment

EdGame uses stealth assessment — students are assessed through natural gameplay without awareness of being tested. This eliminates test anxiety and captures authentic behavior. The assessment is "invisible" because:

- Questions are gated behind game mechanics (answer to attack, answer to build)
- Strategic choices reveal problem-solving approaches
- Social interactions reveal collaboration skills
- Response patterns reveal engagement and affect

### Bayesian Knowledge Tracing (BKT)

BKT estimates the probability that a student has mastered a specific knowledge component, updated after each observation:

- **P(L₀):** Prior probability of mastery (initial estimate)
- **P(T):** Probability of learning (transitioning from unmastered to mastered)
- **P(G):** Probability of guessing correctly despite not having mastered the KC
- **P(S):** Probability of slipping (answering incorrectly despite mastery)

After each student response, the posterior mastery probability is updated using Bayes' theorem.

## Six Analytics Dimensions

EdGame organizes its 50+ metrics across six complementary dimensions:

### D1: Cognitive Knowledge
Measures what students know and can do — factual recall, conceptual understanding, and procedural fluency. Primary metrics: accuracy rate, mastery probability (BKT), knowledge component coverage, misconception identification.

### D2: Behavioral Engagement
Measures how students interact with the learning environment — participation intensity, persistence, and attention patterns. Primary metrics: actions per minute, session duration, voluntary continuation rate, inactivity periods, response time consistency.

### D3: Strategic Behavior & Agency
Measures how students approach problems — planning, resource allocation, experimentation, and self-correction. Primary metrics: strategy diversity index, resource efficiency, exploration rate, error recovery patterns, hint usage patterns.

### D4: Social & Collaborative
Measures how students work together — communication, role fulfillment, knowledge sharing, and team contribution balance. Primary metrics: communication frequency, role adoption consistency, team contribution equity (Gini coefficient), peer teaching instances.

### D5: Affective & Social-Emotional Learning
Measures emotional and motivational states — frustration tolerance, growth mindset indicators, empathy, and emotional regulation. Primary metrics: frustration index (rapid incorrect answers), persistence after failure, emotional response patterns, empathy choices in RPG scenarios.

### D6: Temporal & Longitudinal
Measures learning trajectories over time — mastery growth rates, retention, spacing effects, and long-term trends. Primary metrics: mastery velocity, retention rate, learning curve shape, spacing effect responsiveness.
"""
    add_markdown_content(doc, intro)

    # Coverage Matrix from JSON
    doc.add_heading("Coverage Matrix", level=2)
    p = doc.add_paragraph()
    _apply_inline_formatting(p, "5-games × 6-dimensions analytics coverage matrix showing which dimensions each game primarily (P), secondarily (S), or minimally (·) covers.")
    doc.add_paragraph()

    try:
        cov_data = json.loads(read_file(COVERAGE_FILE))
        dims = cov_data.get("dimensions", [])
        games = cov_data.get("games", [])

        if dims and games:
            # Build table
            headers = ["Game"] + [d["id"] + ": " + d["name"] for d in dims]
            num_cols = len(headers)
            table = doc.add_table(rows=1 + len(games), cols=num_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            for ci, h in enumerate(headers):
                cell = table.cell(0, ci)
                cell.text = h
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
                cell._tc.get_or_add_tcPr().append(shading)

            # Data rows
            for ri, game in enumerate(games):
                table.cell(ri + 1, 0).text = game.get("name", "")
                coverage = game.get("coverage", {})
                for di, dim in enumerate(dims):
                    val = coverage.get(dim["id"], "·")
                    table.cell(ri + 1, di + 1).text = val
                    # Style
                    for p in table.cell(ri + 1, di + 1).paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.size = Pt(10)

            doc.add_paragraph()
    except Exception as e:
        p = doc.add_paragraph()
        p.add_run(f"[Error loading coverage matrix: {e}]").font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # Metric Mapping from JSON
    doc.add_heading("Metric Mapping — Pulse Realms", level=2)
    p = doc.add_paragraph()
    _apply_inline_formatting(p, "Maps core metrics to Pulse Realms telemetry events and computed fields. Source: metric-mapping.json")
    doc.add_paragraph()

    try:
        metric_data = json.loads(read_file(METRIC_FILE))
        dimensions = metric_data.get("dimensions", [])

        for dim in dimensions:
            doc.add_heading(f'{dim.get("id", "")}: {dim.get("name", "")}', level=3)
            metrics = dim.get("metrics", [])
            if metrics:
                headers = ["Metric", "Source Events", "Computation"]
                table = doc.add_table(rows=1 + len(metrics), cols=3)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for ci, h in enumerate(headers):
                    cell = table.cell(0, ci)
                    cell.text = h
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

                for mi, m in enumerate(metrics):
                    table.cell(mi + 1, 0).text = m.get("name", "")
                    sources = m.get("sourceEvents", m.get("source_events", []))
                    if isinstance(sources, list):
                        table.cell(mi + 1, 1).text = ", ".join(sources)
                    else:
                        table.cell(mi + 1, 1).text = str(sources)
                    table.cell(mi + 1, 2).text = m.get("computation", m.get("formula", ""))
                    # Font size
                    for ci in range(3):
                        for p in table.cell(mi + 1, ci).paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)

                doc.add_paragraph()
    except Exception as e:
        p = doc.add_paragraph()
        p.add_run(f"[Error loading metric mapping: {e}]").font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTIONS 12-16: ECD MODELS PER GAME
# ══════════════════════════════════════════════════════════════════════

def write_ecd_game(doc, section_num, game_file, game_name):
    """Write an ECD game section from a markdown file."""
    doc.add_heading(f"{section_num}. ECD Model: {game_name}", level=1)

    content = read_file(game_file)
    # Skip top-level heading
    content = re.sub(r'^#\s+.*\n', '', content, count=1)
    add_markdown_content(doc, content, heading_offset=1)
    add_page_break(doc)


def write_ecd_pulse_realms(doc):
    write_ecd_game(doc, 12, ECD_DIR / "pulse-realms.md", "Pulse Realms — Team Arena")

def write_ecd_concept_cascade(doc):
    write_ecd_game(doc, 13, ECD_DIR / "concept-cascade.md", "Concept Cascade — Tower Defense")

def write_ecd_lab_explorer(doc):
    write_ecd_game(doc, 14, ECD_DIR / "lab-explorer.md", "Lab Explorer — Virtual Science Lab")

def write_ecd_knowledge_quest(doc):
    write_ecd_game(doc, 15, ECD_DIR / "knowledge-quest.md", "Knowledge Quest — Turn-Based RPG")

def write_ecd_survival_equation(doc):
    write_ecd_game(doc, 16, ECD_DIR / "survival-equation.md", "Survival Equation — Collaborative Puzzle Survival")


# ══════════════════════════════════════════════════════════════════════
# SECTION 17: EDGAME ANALYTICS BLUEPRINT
# ══════════════════════════════════════════════════════════════════════

def write_blueprint_section(doc):
    doc.add_heading("17. EdGame Analytics Blueprint (Research Study)", level=1)

    p = doc.add_paragraph()
    _apply_inline_formatting(p, "This section contains the complete EdGame Analytics Blueprint — a comprehensive research study covering the foundational assessment framework, metrics taxonomy, game design principles, dashboard specifications, competitive landscape analysis, implementation roadmap, and ethical considerations.")
    doc.add_paragraph()

    content = read_file(BLUEPRINT_FILE)
    # Skip top-level heading since we have our own
    content = re.sub(r'^#\s+.*\n', '', content, count=1)
    add_markdown_content(doc, content, heading_offset=1)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 18: IMPLEMENTATION STATUS
# ══════════════════════════════════════════════════════════════════════

def write_implementation_status(doc):
    doc.add_heading("18. Implementation Status", level=1)

    content = """\
## Sprint 1 Deliverables (Completed)

- Monorepo scaffold: pnpm + Turborepo with Next.js 14, shared TypeScript packages
- SurrealDB schema: 12 tables with fields, permissions, and seed data
- Teacher dashboard UI: class overview, student detail, session timeline views
- Phase 1 API routes: auth, classes, sessions, telemetry, analytics endpoints
- Evidence-Centered Design mappings for all 5 games
- Analytics coverage matrix and metric mapping for Pulse Realms
- Architecture Decision Record (ADR-001) with 6 binding decisions

## Sprint 2 Deliverables (Completed)

- Pulse Realms game migration into monorepo with backend integration
- Telemetry event pipeline: client → API → SurrealDB
- Real-time dashboard data flow
- Comprehensive documentation consolidation (this document)
- System architecture and algorithm diagrams (7 HTML visualizations)

## Sprint 3 Targets (In Progress)

- BKT mastery estimation engine implementation
- Adaptive difficulty system (FSM-based, connected to BKT)
- Question engine with KC-aware question selection
- Dashboard enhancements: real-time metric visualization, class comparison views
- AI question generation pipeline (Stage 1-2: upload + KC extraction)
- SpacetimeDB evaluation and prototype for Pulse Realms multiplayer
- Tablet adaptation feasibility study for Pulse Realms
"""
    add_markdown_content(doc, content)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# SECTION 19: REPOSITORY STRUCTURE
# ══════════════════════════════════════════════════════════════════════

def write_repo_structure(doc):
    doc.add_heading("19. Repository Structure Overview", level=1)

    content = """\
The EdGame repository is organized as a monorepo using pnpm workspaces and Turborepo for build orchestration.

```
edgame/
├── apps/
│   ├── web/                    # Next.js 14 teacher/student dashboard
│   └── pulse-realms/           # KAPLAY.js 3v3 team arena game
├── packages/
│   ├── shared-types/           # TypeScript interfaces (shared across apps)
│   ├── analytics-engine/       # Metric computation & BKT engine
│   └── ui/                     # Shared React components
├── docs/
│   ├── adr/                    # Architecture Decision Records
│   ├── architecture/           # Data model, API routes, monorepo structure
│   └── assessment/             # ECD mappings, coverage matrix, metrics
├── TIE204Assignments/
│   └── diagrams/               # System architecture & algorithm diagrams
├── turbo.json                  # Turborepo pipeline configuration
├── pnpm-workspace.yaml         # Workspace package definitions
└── package.json                # Root package with shared scripts
```

**Key conventions:**
- All TypeScript interfaces live in `packages/shared-types` and are imported by both apps and API routes
- Game telemetry events follow the schema defined in `docs/assessment/metric-mapping.json`
- Database migrations are version-controlled in `docs/architecture/data-model.surql`
- Each game has its own ECD mapping document in `docs/assessment/ecd/`
"""
    add_markdown_content(doc, content)


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    print("Creating EdGame Living Document...")
    doc = Document()
    setup_styles(doc)

    # Build document section by section
    sections = [
        ("Title Page", write_title_page),
        ("Executive Summary", write_executive_summary),
        ("System Overview", write_system_overview),
        ("ADR", write_adr_section),
        ("Architecture & Algorithms", write_architecture_algorithms),
        ("SpacetimeDB", write_spacetimedb_section),
        ("Question Generation", write_question_generation_section),
        ("Tablet vs PC", write_tablet_vs_pc_section),
        ("Data Model", write_data_model),
        ("API Spec", write_api_spec),
        ("Monorepo Structure", write_monorepo_structure),
        ("Assessment Framework", write_assessment_framework),
        ("ECD: Pulse Realms", write_ecd_pulse_realms),
        ("ECD: Concept Cascade", write_ecd_concept_cascade),
        ("ECD: Lab Explorer", write_ecd_lab_explorer),
        ("ECD: Knowledge Quest", write_ecd_knowledge_quest),
        ("ECD: Survival Equation", write_ecd_survival_equation),
        ("Analytics Blueprint", write_blueprint_section),
        ("Implementation Status", write_implementation_status),
        ("Repository Structure", write_repo_structure),
    ]

    for name, func in sections:
        print(f"  Writing: {name}...")
        func(doc)

    print(f"Saving to {OUTPUT_FILE}...")
    doc.save(str(OUTPUT_FILE))
    size_kb = OUTPUT_FILE.stat().st_size / 1024
    print(f"Done! File size: {size_kb:.0f} KB")


if __name__ == "__main__":
    main()
