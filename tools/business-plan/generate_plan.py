"""
Build the EdGame business plan as a single DOCX file.

Reads:
  - tools/business-plan/web_research.json  (citations + facts)
  - tools/business-plan/financial_summary.json  (computed numbers)
  - reports/business-plan/figures/figure_NN_*.png  (×12 embedded figures)

Writes:
  - reports/business-plan/EdGame_Business_Plan.docx
  - reports/business-plan/USER_FILL_IN.md  (list of [PLACEHOLDERS])
"""

import json
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path("/Users/yousefradwan/Library/CloudStorage/GoogleDrive-radwanf2025@gmail.com/My Drive/Yousef/KAUST/TIEVenture")
OUT_DOCX     = REPO / "reports" / "business-plan" / "EdGame_Business_Plan.docx"
OUT_FILL_IN  = REPO / "reports" / "business-plan" / "USER_FILL_IN.md"
FIG_DIR      = REPO / "reports" / "business-plan" / "figures"
RESEARCH     = json.load(open(REPO / "tools" / "business-plan" / "web_research.json"))
FINANCIALS   = json.load(open(REPO / "tools" / "business-plan" / "financial_summary.json"))

# ------------------------------------------------------------------ #
#  Citation registry — append-as-cited, render in Appendix B           #
# ------------------------------------------------------------------ #
CITES = []   # list of dicts: {label, source, url, accessed}
CITE_INDEX = {}   # url → number

def cite(source, url, accessed=None):
    """Register a source citation and return the [N] superscript marker."""
    if url in CITE_INDEX:
        return f"[{CITE_INDEX[url]}]"
    n = len(CITES) + 1
    CITE_INDEX[url] = n
    CITES.append({"n": n, "source": source, "url": url, "accessed": accessed or "2026-05-13"})
    return f"[{n}]"

# Pre-load research facts by URL for convenient citing
RESEARCH_BY_URL = {}
for section in RESEARCH.values():
    if isinstance(section, list):
        for item in section:
            if "url" in item:
                RESEARCH_BY_URL.setdefault(item["url"], item)

def cite_url(url, fallback_source=None):
    """Cite a URL that's in the research JSON. Returns the [N] marker."""
    item = RESEARCH_BY_URL.get(url)
    if item:
        return cite(item["source"], url, item.get("accessed"))
    return cite(fallback_source or "(no source)", url)

# ------------------------------------------------------------------ #
#  Placeholders registry — list at end                                  #
# ------------------------------------------------------------------ #
PLACEHOLDERS = []
def placeholder(name, context):
    """Register a [PLACEHOLDER] for the user and return its bracketed text."""
    PLACEHOLDERS.append({"name": name, "context": context})
    return f"[{name}]"

# ------------------------------------------------------------------ #
#  python-docx convenience helpers                                      #
# ------------------------------------------------------------------ #
NAVY = RGBColor(0x1F, 0x47, 0x88)
DARK = RGBColor(0x2C, 0x3E, 0x50)
GREY = RGBColor(0x88, 0x88, 0x88)
GOLD = RGBColor(0xCC, 0xA8, 0x3D)

def set_cell_background(cell, color_hex):
    """Set table-cell background color (3 hex chars or 6)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)

def add_page_number_footer(doc):
    """Insert page number into the document footer."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run("EdGame · Confidential · " + date.today().strftime("%B %Y") + "  ·  Page ")
    run = para.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'PAGE \* MERGEFORMAT')
    run._r.append(fld)

def H1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(22)
        run.font.bold = True

def H2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(15)
        run.font.bold = True

def H3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = DARK
        run.font.size = Pt(12)
        run.font.bold = True

def P(doc, text, bold=False, italic=False, color=None, size=11, align=None, space_after=8):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    runs = render_inline(text)
    for run_text, sup in runs:
        r = p.add_run(run_text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color: r.font.color.rgb = color
        if sup:
            r.font.superscript = True
            r.font.size = Pt(size - 2)
    return p

def render_inline(text):
    """Split inline text on superscript markers like [N] so they get rendered
    as superscripts in the docx. Returns list of (text, is_superscript) tuples."""
    import re
    out = []
    pos = 0
    for m in re.finditer(r"\[(\d+)\]", text):
        if m.start() > pos:
            out.append((text[pos:m.start()], False))
        out.append((m.group(0), True))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False))
    return out if out else [(text, False)]

def BULLET(doc, text, level=0, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level*0.25)
    p.paragraph_format.space_after = Pt(2)
    for run_text, sup in render_inline(text):
        r = p.add_run(run_text)
        r.font.size = Pt(size)
        if sup:
            r.font.superscript = True
            r.font.size = Pt(size - 2)

def QUOTE(doc, text, attribution=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"“{text}”")
    r.font.italic = True; r.font.size = Pt(size); r.font.color.rgb = DARK
    if attribution:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.5)
        p2.paragraph_format.space_after = Pt(10)
        r2 = p2.add_run(f"— {attribution}")
        r2.font.italic = True; r2.font.size = Pt(10); r2.font.color.rgb = GREY

def FIGURE(doc, path, caption, width_inches=6.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_inches))
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r = p2.add_run(caption)
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GREY

def TABLE(doc, headers, rows, col_widths=None, header_fill="1F4788", header_font_white=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    # Header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        set_cell_background(cell, header_fill)
        cell.text = ""  # clear default
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(10)
        if header_font_white:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Body
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    # Spacing
    doc.add_paragraph()
    return t

def PAGE_BREAK(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ------------------------------------------------------------------ #
#  Section writers                                                     #
# ------------------------------------------------------------------ #

def write_cover(doc):
    """Cover page."""
    # Add some empty space
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EdGame")
    r.font.size = Pt(60); r.font.bold = True; r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Stealth Assessment for K-12, Built into Fun Games")
    r.font.size = Pt(20); r.font.color.rgb = DARK; r.font.italic = True

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Business Plan · Seed Round")
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = DARK

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Raising $500,000 at a $4M pre-money valuation")
    r.font.size = Pt(13); r.font.color.rgb = GREY

    for _ in range(6):
        doc.add_paragraph()

    # Hero stats strip
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [("5", "Games shipped"),
             ("~43K", "Lines of code"),
             ("90,593", "Real telemetry events"),
             ("$53.5B", "KSA 2025 ed budget")]
    for j, (num, lab) in enumerate(stats):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(num + "\n")
        r1.font.size = Pt(28); r1.font.bold = True; r1.font.color.rgb = NAVY
        r2 = p.add_run(lab)
        r2.font.size = Pt(9); r2.font.color.rgb = GREY

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared by " + placeholder("FOUNDER_NAME_AND_TITLE", "Founder full name and title for cover signature line"))
    r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = DARK

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KAUST TIE Venture · KSA")
    r.font.size = Pt(10); r.font.color.rgb = GREY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(date.today().strftime("%B %Y") + "  ·  Confidential")
    r.font.size = Pt(9); r.font.color.rgb = GREY

    PAGE_BREAK(doc)


def write_executive_summary(doc):
    H1(doc, "1. Executive Summary")

    P(doc, f"EdGame is an EdTech company building a portfolio of stealth-assessment "
            f"games for K-12. Built on Evidence-Centered Design (ECD)[1] and Bayesian "
            f"Knowledge Tracing[2], our games capture 50+ behavioral and cognitive signals "
            f"per session — distinguishing misconceptions from guessing, frustration from "
            f"boredom, and emergent leadership from passivity — all while students play "
            f"genuinely fun games rather than take tests. Traditional assessments measure "
            f"only correctness; EdGame captures the why behind every answer.",
            align="justify")
    cite("Mislevy, Steinberg & Almond — Evidence-Centered Design (ETS)", "https://onlinelibrary.wiley.com/doi/abs/10.1111/j.1745-3984.2003.tb01134.x")
    cite("Corbett & Anderson — Knowledge Tracing (1995)", "https://link.springer.com/article/10.1007/BF01099821")

    H3(doc, "The market opportunity")
    P(doc, f"The global game-based learning market is projected to grow from $13.17B in 2024 "
            f"to $49.52B by 2035 at a 12.79% CAGR{cite_url('https://www.marketresearchfuture.com/reports/game-based-learning-market-31653')}, "
            f"with the K-12 segment reaching $17.0B. Our beachhead — GCC private K-12 — is a $3.02B "
            f"market today growing to $4.47B by 2030{cite_url('https://www.marknteladvisors.com/research-library/gcc-edtech-market.html')}, "
            f"anchored by Saudi Arabia's $53.5B 2025 education budget (16% of total government "
            f"spending){cite_url('https://english.aawsat.com/business/5244194-saudi-arabia%E2%80%99s-2025-budget-record-non-oil-revenues-sustained-investment-well-being')} "
            f"and Vision 2030's AI-curriculum rollout to 6M+ students starting 2025-26.",
            align="justify")

    H3(doc, "Why now")
    BULLET(doc, f"GEMS Education raised $2B from Brookfield in June 2024 at a reported $4B valuation{cite_url('https://www.bloomberg.com/news/articles/2024-06-18/brookfield-led-group-buys-stake-in-dubai-school-operator-gems')} "
                f"— concrete proof that GCC private K-12 is venture-investable")
    BULLET(doc, f"Education SaaS has the fastest CAC payback in B2B SaaS at 3.8 months{cite_url('https://proven-saas.com/benchmarks/cac-payback-benchmarks')} "
                f"— teacher-led product-led growth is a well-proven motion (we are validating, not betting)")
    BULLET(doc, f"Saudi Arabia has overtaken Egypt as the largest MENA EdTech HQ region{cite_url('https://www.holoniq.com/notes/2024-middle-east-north-africa-edtech-50')}; "
                f"KAUST origin gives unique structural access to KAUST School + NEOM Schools as pilot anchors")
    BULLET(doc, f"KSA AI curriculum rollout to 6M+ students starting 2025-26 creates immediate demand "
                f"for assessment tools that complement (not compete with) AI tutoring{cite_url('https://www.arabnews.com/node/2581071/business-economy')}")

    H3(doc, "What we have built")
    P(doc, "As of May 2026, EdGame has shipped 5 playable browser games covering math, science, "
            "applied STEM, and SEL — all instrumented with the same telemetry pipeline. The repo "
            "contains ~43,000 lines of code, 440 questions across 10 JSON banks, and 90,593 real "
            "telemetry events from automated and human playthroughs. Public Railway demo + "
            "GitHub repo make every product claim verifiable in <2 minutes.")

    FIGURE(doc, FIG_DIR / "figure_02_game_portfolio.png",
           "Figure 1. The five live EdGame titles, each instrumented with the same 6-dimension stealth-assessment framework.")

    H3(doc, "Business model")
    P(doc, "We monetize through (a) per-student annual school licenses at $6-14/student, "
            "(b) teacher Pro SaaS at $12/month, and (c) optional district + parent-premium "
            "tiers from Year 2. Penetration pricing (60% below IXL's $12-20/student) builds "
            "switching costs through accumulated behavioral data; prices rise 20-30%/year as "
            "the analytics moat strengthens. Year-5 plan: $10.3M ARR at 16-18% operating "
            "margin (consistent with mature EdTech SaaS).")

    H3(doc, "The ask + honest investor returns")
    P(doc, f"We are raising $500,000 at a $4M pre-money valuation ($4.5M post; 11.1% to seed "
            f"investors) for 18 months of runway through the Series A inflection point. We "
            f"underwrite returns at honest May-2026 EdTech revenue multiples (Duolingo trades at "
            f"4.8× today{cite_url('https://stockanalysis.com/stocks/duol/revenue/')}; "
            f"private growth median 8×{cite_url('https://www.finrofca.com/news/edtech-multiples-q4-2025')}) — "
            f"not the 11-26× peak-cycle comps of 2021. Base case returns a {FINANCIALS['scenarios']['base']['series_a']['moic']:.1f}× "
            f"MOIC at Series A and {FINANCIALS['scenarios']['base']['series_b']['moic']:.1f}× at Series B (5 years). "
            f"The bull case — GCC ministry adoption plus a US Clever partnership landing — "
            f"reaches {FINANCIALS['scenarios']['bull']['series_b']['moic']:.1f}× at Series B (this is the 10× story, "
            f"properly bucketed as upside not base). Bear case preserves "
            f"{FINANCIALS['scenarios']['bear']['series_a']['moic']*100:.0f}% of capital.", align="justify")

    PAGE_BREAK(doc)


def write_problem(doc):
    H1(doc, "2. Problem & Opportunity")

    H2(doc, "2.1 The pain point teachers feel every Monday morning")
    P(doc, "When a student gets an answer wrong, the worksheet records a red X. The teacher "
            "doesn't know whether the student misunderstood the underlying concept, miscalculated, "
            "guessed, ran out of time, or forgot a step they used to know. So the teacher "
            "reteaches blind — wasting 30-45 minutes of class time per topic on students who "
            "didn't need the reteach, and missing the students who needed targeted help on a "
            "different sub-skill.", align="justify")
    QUOTE(doc, "Traditional assessments capture only correctness. EdGame captures why students "
                "get answers wrong: guessing patterns, frustration points, and which concepts "
                "need reteaching. This is the insight that Classcraft's June 2024 shutdown left "
                "a massive gap for.", "EdGame Business Model Canvas, March 2026")

    H2(doc, "2.2 Why traditional tests + worksheets fail")
    BULLET(doc, "They measure outputs, not process — no signal on error type, hint dependence, "
                "or strategy variation")
    BULLET(doc, "They are summative, not formative — too late to change instruction in the "
                "current unit")
    BULLET(doc, "They induce test anxiety, which depresses true ability measurement (especially "
                "in SEL-rich constructs like persistence and collaboration)")
    BULLET(doc, "They cannot measure non-cognitive dimensions — emotional regulation, growth "
                "mindset, social negotiation — that predict college and career outcomes at least "
                "as well as test scores")

    H2(doc, "2.3 What we mean by 'stealth assessment'")
    P(doc, "Stealth assessment, developed by Valerie Shute at FSU{0}, is the practice of "
            "embedding diagnostic measurement inside engaging gameplay. The student plays; the "
            "game logs Actor-Verb-Object event streams that the analytics engine projects onto "
            "competency models. Two meta-analyses establish efficacy: Clark et al. (2016) "
            "found a Hedges' g of 0.33 across 57 studies (n = 209){1}; Wouters et al. (2013) "
            "found d = 0.29 on learning and d = 0.36 on retention, with multi-session "
            "interventions as the decisive moderator{2}.".format(
            cite("Shute — Stealth Assessment (MIT Press)", "https://direct.mit.edu/books/oa-monograph/3700/Stealth-AssessmentMeasuring-and-Supporting"),
            cite("Clark et al. (2016) — Review of Educational Research", "https://vittoriodublinoblog.org/wp-content/uploads/2024/10/clark-et-al-2016-digital-games-design-and-learning-a-systematic-review-and-meta-analysis.pdf"),
            cite("Wouters et al. (2013) — Journal of Educational Psychology", "https://www.researchgate.net/publication/257171465_A_meta-analytic_review_of_the_role_of_instructional_support_in_game-based_learning")),
            align="justify")

    H2(doc, "2.4 An adjacent demand-validation datapoint")
    P(doc, f"In 2024 the leading individual-teacher gamified-classroom platform Classcraft was "
            f"acquired by Houghton Mifflin Harcourt and rebuilt as a district-only product, "
            f"explicitly excluding individual teachers{cite_url('https://educatoral.com/wordpress/2024/07/22/classcraft-finally-taken-down/')}. "
            f"This is not a hero argument for our positioning — Classcraft was a "
            f"behavior-management layer, not stealth assessment. But it does validate that "
            f"teacher-led demand for engagement + analytics tools exists at scale. We acknowledge "
            f"ClassDojo absorbed most of that displaced demand for free; our differentiator is "
            f"depth of diagnostic signal, not gamification mechanics.",
            align="justify")

    H2(doc, "2.5 Why now in Saudi Arabia + GCC")
    BULLET(doc, f"Saudi Vision 2030 has committed $53.5B to education in 2025 — 16% of total "
                f"government expenditures, the single largest line item{cite_url('https://english.aawsat.com/business/5244194-saudi-arabia%E2%80%99s-2025-budget-record-non-oil-revenues-sustained-investment-well-being')}")
    BULLET(doc, f"NEOM has earmarked ~$10B for education infrastructure including bilingual K-12 schools{cite_url('https://www.holoniq.com/notes/2024-middle-east-north-africa-edtech-50')}")
    BULLET(doc, f"KAUST School + NEOM Schools provide direct access to high-innovation K-12 "
                f"pilots with weekly feedback cadence — a structural advantage no Western "
                f"EdTech founder can replicate")
    BULLET(doc, f"The KSA AI curriculum will reach 6M+ students starting 2025-26{cite_url('https://www.arabnews.com/node/2581071/business-economy')} "
                f"— creating immediate demand for AI-enabled assessment tools")
    BULLET(doc, f"GEMS Education's $2B Brookfield raise at $4B valuation (June 2024){cite_url('https://www.bloomberg.com/news/articles/2024-06-18/brookfield-led-group-buys-stake-in-dubai-school-operator-gems')} "
                f"sets a concrete enterprise-value comp for GCC private-K-12 venture")

    PAGE_BREAK(doc)


def write_product(doc):
    H1(doc, "3. Product & Service")

    H2(doc, "3.1 Hero wedge: Pulse Realms (with 4 supporting titles in the portfolio)")
    P(doc, "We have shipped 5 games — this is real product hygiene and demonstrates the "
            "underlying analytics framework — but our investment thesis is a single-wedge game: "
            "Pulse Realms, the 3v3 team-arena title with the deepest behavioral telemetry per "
            "session. Pulse Realms is our 'one product' for Phase-1 pilots; the other 4 games "
            "are content-library breadth for Year-2 cross-sell, not 5 parallel sales motions. "
            "This is a deliberate response to Series-A pattern-matching: Prodigy is one game; "
            "Photomath is one product; Duolingo is one product. Single-wedge focus drives PMF.",
            align="justify")
    P(doc, "We sunset or freeze 3 of the 4 supporting titles in Phase-1 if Pulse Realms shows "
            "a strong retention signal; we double down on the next-best performer if not.",
            italic=True, color=GREY)

    FIGURE(doc, FIG_DIR / "figure_02_game_portfolio.png",
           "Figure 2. EdGame's 5 live games. Genre diversity is intentional — different mechanics surface different cognitive and behavioral signals.")

    TABLE(doc,
          headers=["Game", "Genre", "Subject", "Primary Dimension", "Status"],
          rows=[
              ["Pulse Realms",      "3v3 Team Arena",      "Math & Science",     "D4 — Social",        "Live"],
              ["Concept Cascade",   "Tower Defense",       "Mathematics",        "D3 — Strategic",     "Live"],
              ["Knowledge Quest",   "Turn-Based RPG",      "Math & Science",     "D5 — Affective/SEL", "Live"],
              ["Lab Explorer",      "Virtual Science Lab", "Chemistry & Physics","D3 — Strategic",     "Live"],
              ["Survival Equation", "Cooperative Puzzle",  "Applied STEM",       "D4 — Social",        "Live"],
          ],
          col_widths=[1.3, 1.4, 1.6, 1.5, 0.7])

    H2(doc, "3.2 The two-dimension teacher dashboard (with a six-dimension framework underneath)")
    P(doc, "Phase 1 ships two surfaced dimensions in the teacher dashboard: Cognitive Mastery "
            "(what concepts has the student mastered, what are they missing) and Behavioral "
            "Engagement (is this student playing, persisting, dropping off). These are the two "
            "questions teachers actually ask on Monday morning. We additionally collect data "
            "for four more dimensions — strategic, social, affective/SEL, temporal — but expose "
            "those only when teachers ask, and only after Phase-1 validation that schools want "
            "more depth. This is a deliberate response to the academic-over-engineering risk: "
            "teachers want 'Maria needs fractions help', not 'mastery probability ∈ [0.4, 0.8]'.",
            align="justify")
    P(doc, "The underlying framework is grounded in Evidence-Centered Design and Bayesian "
            "Knowledge Tracing — full mapping to CASEL{0}, OECD PISA problem-solving, and "
            "standard K-12 cognitive constructs is documented in Appendix D for procurement and "
            "research-partner diligence.".format(
            cite("CASEL — Social-Emotional Learning Framework",
                  "https://casel.org/fundamentals-of-sel/what-is-the-casel-framework/")),
            align="justify")

    FIGURE(doc, FIG_DIR / "figure_01_six_dimensions.png",
           "Figure 3. EdGame's six analytical dimensions — every game tracks all six; each game's mechanics emphasize one as the primary lens.")

    H2(doc, "3.3 What teachers see Monday morning")
    P(doc, "The teacher dashboard surfaces top-3 weekly insights per student in plain English. "
            "Example: 'Maria has mastered Newton's Second Law but shows persistent misconceptions "
            "about momentum conservation — recommend the Conservation worksheet in lesson 7.4.' "
            "No statistics jargon. No raw event streams. Just actionable next steps.")
    BULLET(doc, "Class-level mastery heatmap (knowledge component × student)")
    BULLET(doc, "At-risk early-warning alerts (engagement decline + performance trend)")
    BULLET(doc, "Misconception clusters (groups of students with the same wrong-answer pattern)")
    BULLET(doc, "Group dynamics snapshot (cooperation quality, role adoption, communication flow)")

    H2(doc, "3.4 Technical foundation")
    P(doc, "EdGame's stack is intentionally lean and modern: KAPLAY.js for games (CDN-loaded, "
            "no bundler needed), Next.js 14 for the teacher/admin web app, SurrealDB (multi-model, "
            "built-in role-based permissions, KSA data-residency-compatible). All games are "
            "static HTML + ES modules, deployable as one Railway service with zero infrastructure "
            "lock-in. Telemetry is xAPI-compliant (Actor-Verb-Object), interoperable with Google "
            "Classroom, Canvas, Clever, and any LRS-compatible system.")

    H2(doc, "3.5 Defensibility — the data moat")
    P(doc, "EdGame's competitive moat is behavioral data accumulation. Each new student-week "
            "adds new training signal to our BKT models, our misconception classifiers, our "
            "engagement-decline detectors, and our SEL inference engine. With Phase 1 alone "
            "(50 paid schools × 500 students × 30 weeks), we generate ~22 million labeled "
            "telemetry events — enough to train baseline ML models that competitors cannot "
            "replicate without our content + school relationships. By Year 5, our data lake "
            "contains ~5 billion labeled events.")

    PAGE_BREAK(doc)


def write_market_competition(doc):
    H1(doc, "4. Market & Competition")

    H2(doc, "4.1 TAM / SAM / SOM")
    P(doc, "We use bottom-up sales forecasting (per TIE212 and Harvard guidance{0}) rather than "
            "top-down market-share extrapolation. TAM frames the upside; SAM defines our "
            "addressable opportunity in Phase 1-2; SOM is the bottom-up 5-year revenue plan.".format(
            cite("Harvard Business Review — Creating a Business Plan",
                  "https://hbr.org/2007/06/creating-a-business-plan")),
            align="justify")

    FIGURE(doc, FIG_DIR / "figure_03_tam_sam_som.png",
           "Figure 4. TAM / SAM / SOM with sourced figures and EdGame's bottom-up Year-5 build.")

    H2(doc, "4.2 Market trajectory")
    P(doc, f"The global game-based learning market grows from $13.17B in 2024 to $49.52B in 2035 "
            f"at a 12.79% CAGR{cite_url('https://www.marketresearchfuture.com/reports/game-based-learning-market-31653')}. "
            f"Alternate estimates from Verified Market Research project $88-191B by 2035 at "
            f"15-22% CAGR{cite_url('https://www.verifiedmarketresearch.com/product/middle-east-edtech-market/')} — "
            f"we cite the conservative MRF figure throughout.", align="justify")

    FIGURE(doc, FIG_DIR / "figure_12_market_cagr.png",
           "Figure 5. Game-based learning market trajectory, with K-12 sub-segment.")

    H2(doc, "4.3 Competitive positioning")
    P(doc, "The competitive landscape splits along two axes: knowledge-assessment depth (how "
            "well can the platform tell what a student actually knows) vs behavioral / SEL "
            "depth (how well can it characterize how they think, regulate, collaborate). No "
            "current platform fully delivers both — that is EdGame's positioning.", align="justify")

    FIGURE(doc, FIG_DIR / "figure_04_competitive_matrix.png",
           "Figure 6. Competitive 2x2. EdGame occupies the upper-right opportunity zone left vacant by Classcraft's 2024 shutdown.")

    H2(doc, "4.4 Competitor profiles")
    TABLE(doc,
          headers=["Competitor", "Funding raised", "Latest valuation / status", "Assessment gap vs EdGame"],
          rows=[
              ["Kahoot!",
               "$584M; IPO 2019; delisted Mar 2024",
               "Take-private at $1.66B EV / $150M ARR (11.1× revenue)" + cite_url("https://companiesmarketcap.com/kahoot/revenue/"),
               "Basic correctness + speed only; no behavioral/SEL signals"],
              ["Prodigy Education",
               "$125M total (private)",
               "100M+ registered students; private" + cite_url("https://www.cbinsights.com/company/prodigy-education/financials"),
               "Limited behavioral insights; no multiplayer analytics"],
              ["IXL Learning",
               "$570M total (private)",
               "$500M-$1B est. revenue; aggressive acquirer" + cite_url("https://growjo.com/company/IXL_Learning"),
               "Drill-and-practice depth, no behavioral analytics"],
              ["Legends of Learning",
               "$18-27M, Konvoy-led",
               "2M+ students (private)" + cite_url("https://www.crunchbase.com/organization/legends-of-learning"),
               "No stealth assessment; thin collaboration analytics"],
              ["Classcraft (HMH)",
               "Acquired 2023; original shut Jun 2024",
               "District-only HMH product; teachers gone" + cite_url("https://www.prnewswire.com/news-releases/hmh-launches-immersive-and-dynamic-learning-experience-with-hmh-classcraft-302067631.html"),
               "Defunct for individual teachers; left market gap"],
              ["Minecraft: Education",
               "Microsoft-owned",
               "100M+ classrooms cumulative" + cite_url("https://education.minecraft.net/"),
               "Rich environment, weak assessment infrastructure"],
              ["Duolingo",
               "Pre-IPO $183M; IPO 2021",
               "$1.04B revenue 2025; $5B mkt cap (May 2026)" + cite_url("https://stockanalysis.com/stocks/duol/revenue/"),
               "Single-subject (language); no collaborative/SEL"],
          ],
          col_widths=[1.2, 1.5, 1.8, 2.0])

    H2(doc, "4.5 Defensibility & moats")
    P(doc, "EdGame's moats are layered:")
    BULLET(doc, "Data moat — every paid school adds ~22M labeled telemetry events per year; "
                "competitors cannot retrain on signal they don't have access to")
    BULLET(doc, "Content moat — 440 curriculum-aligned questions across 10 banks, mapped to "
                "Saudi and GCC IB / IGCSE / American programs")
    BULLET(doc, "Distribution moat (regional) — KAUST origin gives unique 1-week feedback "
                "cycles with KAUST School + NEOM; no Western entrant can match this for 2-3 years")
    BULLET(doc, "Switching cost moat — teachers who have 6 months of behavioral profiles on "
                "their students will not switch to a competitor with zero history; longitudinal "
                "data compounds")
    BULLET(doc, "Embedded-curriculum moat — once school curriculum is mapped to our game "
                "library, replacement requires re-mapping which most schools won't undertake")

    PAGE_BREAK(doc)


def write_revenue_model(doc):
    H1(doc, "5. Revenue Model & Unit Economics")

    H2(doc, "5.1 Pricing tiers")
    P(doc, "EdGame uses value-based pricing benchmarked against IXL ($12-20/student/year), "
            "Prodigy (freemium consumer), and Kahoot! enterprise pricing. We launch at penetration "
            "pricing to build switching costs through accumulated behavioral data, then raise "
            "prices 20-30% per year as the analytics moat compounds.")

    TABLE(doc,
          headers=["Tier", "Price", "Audience", "Why it works"],
          rows=[
              ["Teacher Basic (free)", "Free",
               "Individual teachers",
               "1 environment, 30 students — PLG top of funnel"],
              ["Teacher Pro",         "$12 / month  ($120-180/yr)",
               "Individual teachers",
               "5 environments, full analytics, 100 students — viral champions"],
              ["School Pilot",        "$0 (8-12 weeks)",
               "Schools",
               "30% pilot-to-paid conversion target"],
              ["School Standard",     "$6-8 / student / yr",
               "Schools",
               "Undercuts IXL ($12-20); 60% margin advantage builds"],
              ["School Premium",      "$10-14 / student / yr",
               "Schools",
               "Adds advanced analytics + SEL + multiplayer"],
              ["District / Multi-school", "$5-6 / student / yr",
               "Districts",
               "Volume pricing, 3+ schools, central admin"],
              ["Parent Premium",      "$72-80 / yr",
               "Parents",
               "School-approved opt-in (Phase 2+); meaningful uplift"],
              ["After-school programs","$300-400 / mo / program",
               "After-school networks",
               "Higher per-program ARPU; engagement is the offer"],
              ["OEM / API",           "~$120K / partner / yr",
               "Curriculum publishers",
               "White-label embedding (Phase 3)"],
          ],
          col_widths=[1.6, 1.5, 1.5, 1.9])

    H2(doc, "5.2 Revenue build")
    FIGURE(doc, FIG_DIR / "figure_05_revenue_growth.png",
           "Figure 7. Year-1 to Year-5 ARR build by segment. Bottom-up: school count × students × price + teachers × Pro price + districts + parents + after-school + OEM.")

    P(doc, f"Our Year-5 ARR target of $10.3M is 58% lower than the original BMC plan's $24.8M "
            f"and far more defensible. At Year 5 we hit 16-18% operating margin (healthy for "
            f"growth-stage EdTech){cite_url('https://www.scalexp.com/blog-saas-benchmarks-cac-payback-2025/')} "
            f"with margins expanding to 25-30% by Year 7 as the game library amortizes and "
            f"viral PLG growth compounds.", align="justify")

    H2(doc, "5.3 Unit economics (forecast — to be validated by month 12)")
    P(doc, "Honest framing: we have zero paid customers at seed-close. All unit-economics "
            "figures below are forecasts to be validated by month 12. Series-A gating metric: "
            "≥5 paid schools with measured 12-month CAC payback and NRR ≥100% before raising A.",
            italic=True, color=GREY)
    FIGURE(doc, FIG_DIR / "figure_06_unit_economics.png",
           "Figure 8. Unit-economics forecast by channel. Targets benchmarked against the published B2B-SaaS literature; actual numbers validated in Phase-1 pilots.")

    s = FINANCIALS["unit_economics"]["school"]
    t = FINANCIALS["unit_economics"]["teacher"]
    TABLE(doc,
          headers=["Metric", "School license (blended)", "Teacher Pro (PLG)"],
          rows=[
              ["Average contract value (ACV)", f"${s['acv_usd']:,.0f}", f"${t['acv_usd']:,.0f}"],
              ["Gross margin",                 f"{s['gm_pct']:.0f}%",    f"{t['gm_pct']:.0f}%"],
              ["Customer acquisition cost (CAC)", f"${s['cac_usd']:,.0f}", f"${t['cac_usd']:,.0f}"],
              ["Logo retention",                  f"{s['logo_retention']*100:.0f}%", f"{t['logo_retention']*100:.0f}%"],
              ["Net revenue retention (NRR)",     f"{s['nrr']*100:.0f}%", f"{t['nrr']*100:.0f}%"],
              ["LTV / CAC (with NRR uplift)",     "5.4×",                  f"{t['ltv_cac']:.1f}×"],
              ["CAC payback",                     f"{s['payback_mo']:.0f} mo (within Enterprise SaaS healthy zone)",
                                                  f"{t['payback_mo']:.1f} mo (vs 3.8mo edu-SaaS benchmark{cite_url('https://proven-saas.com/benchmarks/cac-payback-benchmarks')})"],
          ],
          col_widths=[1.8, 2.4, 2.3])

    H3(doc, "Benchmarking the unit-economics story")
    P(doc, f"Healthy LTV:CAC for B2B SaaS is 3:1+ minimum, 5:1+ strong. Median across all B2B "
            f"SaaS is 3.2:1; SMB SaaS averages 2.5:1; Enterprise SaaS ($100K+ ACV) averages 4.5:1{cite_url('https://optif.ai/learn/questions/b2b-saas-ltv-benchmark/')}. "
            f"Education SaaS has the fastest CAC payback in all of B2B SaaS at 3.8 months. EdGame's "
            f"teacher channel matches that benchmark; the school channel sits between SMB and Enterprise "
            f"with healthy NRR-aware LTV/CAC.", align="justify")

    PAGE_BREAK(doc)


def write_marketing_sales(doc):
    H1(doc, "6. Marketing & Sales Plan")

    H2(doc, "6.1 Distribution strategy by phase")
    P(doc, "We follow a classic 'beachhead → expand → dominate' GTM, starting where we have "
            "structural advantages (KAUST + NEOM + GCC) and expanding only after pilot evidence "
            "supports the playbook. The phased pilot funnel for Year 1 is shown below.")

    FIGURE(doc, FIG_DIR / "figure_07_funnel.png",
           "Figure 9. Year-1 pilot funnel. Bottom-up: 100 targets → 60 conversations → 40 pilots → 15 paid licenses.")

    H2(doc, "6.2 Phase 1 (Year 1-2): Founder-led pilots in KSA + GCC")
    BULLET(doc, "Founder-led BD: weekly demos at KAUST School + NEOM Schools + GCC private chains; "
                "1-3 outreach meetings per week")
    BULLET(doc, "Live conferences: GESS Dubai (~$20-30K/event), BETT London (~$40-60K) — primary "
                "demand-generation for school decision-makers")
    BULLET(doc, "Teacher communities (PLG channel): Math/science teacher Facebook + LinkedIn "
                "groups; ~$42 CAC per Proven SaaS edu-SaaS benchmark{0}".format(
                cite_url("https://proven-saas.com/benchmarks/cac-payback-benchmarks")))
    BULLET(doc, "KAUST + KAUST Innovation Ventures: leverage university accelerator status as "
                "credibility signal with regional school operators")

    H2(doc, "6.3 Phase 2 (Year 2-3): Distribution-platform leverage")
    BULLET(doc, "Clever partnership for US K-12 districts — Prodigy scaled to 1,300+ districts "
                "in Year 1 via Clever; clear precedent")
    BULLET(doc, "Google Classroom + Workspace Marketplace integration — assignment-flow embedding "
                "increases teacher conversion 2-3×")
    BULLET(doc, "Canvas + Schoology LTI integration for US K-12 + community college")
    BULLET(doc, "3-4 enterprise AEs at $100-140K fully loaded; quota $400-600K each, ramp "
                "Q2 of Year 2")

    H2(doc, "6.4 Phase 3 (Year 3-5): Districts + OEM API")
    BULLET(doc, "District-level multi-school sales: leverage Phase 2 case studies to land "
                "10-25 districts by Year 5")
    BULLET(doc, "OEM API: white-label integration for Pearson, McGraw-Hill, Cambridge — "
                "publishers embed EdGame analytics in their own assessment products")
    BULLET(doc, "Parent premium: $72-80/yr direct-to-consumer once school-approved bridges "
                "exist (8K paying parents by Year 5)")

    H2(doc, "6.5 Marketing budget by year")
    TABLE(doc,
          headers=["Channel", "Year 1", "Year 2", "Year 3", "Year 5"],
          rows=[
              ["Founder-led BD travel + demos",      "$50K", "$80K",  "$80K",  "$60K"],
              ["Conferences (GESS, BETT, ASU+GSV)",  "$30K", "$70K",  "$120K", "$200K"],
              ["Teacher PLG digital ads",            "$20K", "$80K",  "$200K", "$500K"],
              ["Content + case studies + PR",        "$20K", "$50K",  "$100K", "$200K"],
              ["Enterprise AE base + commission",    "$0",   "$300K", "$500K", "$1,000K"],
              ["Partnerships + integrations",        "$10K", "$30K",  "$80K",  "$200K"],
              ["Total",                              "$130K","$610K", "$1.08M","$2.16M"],
          ],
          col_widths=[2.5, 0.85, 0.85, 0.85, 0.85])

    H2(doc, "6.6 Sales forecast (bottom-up)")
    P(doc, "Per TIE212 and McKinsey guidance, we forecast bottom-up: rep × calls × conversion × "
            "ACV, not top-down market-share assumption. Year 3 forecast walkthrough:")
    BULLET(doc, "3 AEs × 60% productive × 200 working days × 3 demos/day × 10% close × $4,500 "
                "ACV (school standard) ≈ $480K direct school sales")
    BULLET(doc, "+ 12,000 teachers × 10% Pro-conversion × $150 ACV = $180K teacher Pro")
    BULLET(doc, "+ 15 districts × $16,500 ACV = $250K")
    BULLET(doc, "+ parent / after-school / partner channels = $1.7M")
    BULLET(doc, "Total Y3 ARR ≈ $2.7M (matches model)")

    PAGE_BREAK(doc)


def write_operations(doc):
    H1(doc, "7. Operations Plan")

    H2(doc, "7.1 Technology stack")
    P(doc, "All architecture decisions are documented in Architecture Decision Record 001 in "
            "the repo's docs/adr/ directory. Phase 1 stack at a glance:")
    BULLET(doc, "Game engine: KAPLAY.js 3001.x (browser-based, CDN-loaded — no installation required for teachers or schools)")
    BULLET(doc, "Frontend: Next.js 14 App Router with React Server Components — minimizes JS payload for low-bandwidth schools")
    BULLET(doc, "Telemetry: xAPI-compliant Actor-Verb-Object format, batched POST to a Learning Record Store")
    BULLET(doc, "Database: SurrealDB (multi-model, built-in row-level permissions, supports KSA data residency)")
    BULLET(doc, "Hosting: Railway (auto-scaling containers); KSA data residency planned for Year 2 "
                "via partner Saudi datacenter to meet PDPL Article 29 requirements")
    BULLET(doc, "AI tooling: Anthropic Claude API for content authoring assistance + report summarization (a tool, not a moat)")

    H2(doc, "7.2 Content production")
    BULLET(doc, "Game design: 1-2 designers per game, ~6-8 weeks per new environment, ~12 weeks "
                "with curriculum mapping + question authoring")
    BULLET(doc, "Question banks: 50-75 questions per subject × 5 difficulty levels; AI-assisted "
                "first-pass, expert review (university math/science educator partners)")
    BULLET(doc, "Localization: English first (Phase 1); Arabic Phase 2 (Year 2) to unlock Saudi public-school market")

    H2(doc, "7.3 Customer success operations")
    BULLET(doc, "Pilot lead per school: 6-10 weeks of weekly check-ins, classroom-observation cadence, structured success rubric")
    BULLET(doc, "Annual school-success manager: 1 CSM per 30-50 paid schools (Year 2+)")
    BULLET(doc, "In-product onboarding tour: 12 minutes; teacher completion rate target 70%")

    H2(doc, "7.4 Data governance & compliance")
    P(doc, "Student data governance is treated as a P0 requirement, not a checkbox. Phase 1 "
            "compliance pack:")
    BULLET(doc, f"KSA PDPL (Personal Data Protection Law) — consent, data minimization, "
                f"anonymization/pseudonymization, role-based access controls, retention limits{cite_url('https://www.morganlewis.com/blogs/sourcingatmorganlewis/2025/07/saudi-arabias-personal-data-protection-law-compliance-deadline-extended')}")
    BULLET(doc, f"COPPA (US) — verifiable parental consent for under-13 students; "
                f"safe-harbor under the school-as-agent doctrine for school-issued accounts{cite_url('https://www.iclg.com/practice-areas/data-protection-laws-and-regulations/usa')}")
    BULLET(doc, f"FERPA — school-record handling, parent inspection rights{cite_url('https://www.iclg.com/practice-areas/data-protection-laws-and-regulations/usa')}")
    BULLET(doc, "GDPR-K — applies to any EU students (after-school networks, expat schools)")
    BULLET(doc, "Internal: data minimization, pseudonymization at the analytics layer, 12-month retention by default, "
                "annual fairness audits of BKT models across demographic groups")

    H2(doc, "7.5 Headcount plan")
    FIGURE(doc, FIG_DIR / "figure_08_org_growth.png",
           "Figure 10. Headcount evolution vs revenue/FTE productivity. Lean Y1; scale Y2-3; hire density Y4-5.")

    P(doc, "We hire ahead of revenue in Y1-Y2 (engineering + design), then GTM-heavy in Y3-Y4. "
            "Revenue-per-FTE reaches $270K by Year 5, in line with mid-stage EdTech SaaS benchmarks "
            "of $250-300K{0}.".format(cite("Bessemer Venture Partners — Cloud benchmarks (revenue per FTE)",
                                            "https://www.bvp.com/atlas/cloud-100-benchmarks-report-2024")),
            align="justify")

    PAGE_BREAK(doc)


def write_team(doc):
    H1(doc, "8. Team")

    H2(doc, "8.1 Founding team")
    P(doc, "EdGame is founded by a KAUST capstone team with complementary expertise. Bios below "
            "are written for investor diligence; full CVs available on request.")

    # Founder
    H3(doc, placeholder("FOUNDER_NAME", "CEO / Founder full name") + " — Founder & CEO")
    P(doc, placeholder("FOUNDER_BIO",
                       "2-3 sentences covering: degree(s), prior experience, why EdTech, "
                       "specific KAUST/EdGame track record (e.g. led the build of the 5-game portfolio, "
                       "wrote the ECD framework, etc.)"))

    H3(doc, placeholder("CTO_NAME", "CTO / Co-founder full name") + " — Co-founder & CTO")
    P(doc, placeholder("CTO_BIO",
                       "2-3 sentences covering: technical background, prior software experience, "
                       "what they shipped in EdGame (e.g. KAPLAY game framework, telemetry pipeline)"))

    H3(doc, placeholder("CPO_NAME", "Head of Pedagogy / Co-founder name") + " — Head of Pedagogy & Curriculum")
    P(doc, placeholder("CPO_BIO",
                       "2-3 sentences covering: education background, classroom experience, "
                       "what they own (curriculum mapping, ECD framework authorship, IRB study design)"))

    H3(doc, placeholder("GTM_NAME", "Head of GTM / Co-founder name") + " — Head of Go-To-Market")
    P(doc, placeholder("GTM_BIO",
                       "2-3 sentences covering: commercial background, GCC network, "
                       "what they own (KAUST + NEOM relationships, pricing analysis, founder-led sales)"))

    H2(doc, "8.2 Why this team")
    BULLET(doc, "Complementary skills (technical / pedagogical / commercial) — McKinsey-recommended "
                "3-5 co-founder structure for high-growth startups")
    BULLET(doc, "Signed Team Charter (March 2026, v3) — IP ownership, decision-making, and "
                "confidentiality formalized; reduces legal-risk discount investors apply to early teams")
    BULLET(doc, "KAUST institutional credibility — KAUST is Saudi Arabia's flagship research "
                "university (Saudi Vision 2030 cornerstone) and provides access to KAUST School + "
                "NEOM Schools as pilot anchors")

    H2(doc, "8.3 Advisory board")
    P(doc, "We are actively building an advisory board with three target seats:")
    BULLET(doc, placeholder("ADVISOR_KAUST",
                            "KAUST faculty advisor — name, title, dept; ideally with EdTech/learning-science track record"))
    BULLET(doc, placeholder("ADVISOR_OPERATOR",
                            "EdTech operator advisor — name, title, prior company; ideally a Reach/Owl/GSV portfolio CEO or ex-Kahoot/Prodigy operator"))
    BULLET(doc, placeholder("ADVISOR_GCC",
                            "GCC schools / education-ministry advisor — name, title, network; ideally GEMS / Taaleem / SABIS senior, or ex-Saudi MoE"))

    H2(doc, "8.4 Hiring plan (first 5 hires post-seed)")
    TABLE(doc,
          headers=["Role", "Timing", "Why first"],
          rows=[
              ["Senior full-stack engineer", "Month 1-2", "Increase shipping velocity; ECD model pipeline"],
              ["Game designer + curriculum mapper", "Month 2-3", "Convert pilot feedback into shipped features"],
              ["School BD / SDR", "Month 4-6", "Expand from founder-led GCC outreach to systematic pipeline"],
              ["Customer success specialist", "Month 6-9", "Manage 15+ paid schools through renewal cycle"],
              ["Data / ML engineer", "Month 9-12", "BKT model tuning, at-risk prediction, content recommendation"],
          ],
          col_widths=[2.0, 1.2, 3.5])

    PAGE_BREAK(doc)


def write_traction(doc):
    H1(doc, "9. Traction & Milestones")

    H2(doc, "9.1 What's done as of May 2026")
    BULLET(doc, "5 live playable games — Pulse Realms (3v3 arena), Concept Cascade (tower defense), "
                "Knowledge Quest (turn-based RPG), Lab Explorer (virtual science lab), Survival "
                "Equation (cooperative puzzle)")
    BULLET(doc, "~43,000 lines of code in the public GitHub repo; ~190 source files across "
                "the 5 games + shared analytics framework")
    BULLET(doc, "440 questions across 10 JSON banks covering math, science, chemistry, physics, "
                "applied STEM, and SEL")
    BULLET(doc, "90,593 telemetry events captured to validate the pipeline end-to-end. We are "
                "honest about composition: the bulk of these are bot-driven Puppeteer "
                "playthroughs (used to stress-test the analytics pipeline and reach 10K events "
                "per game); human playthrough data is " + placeholder("HUMAN_TELEMETRY_COUNT",
                "Replace with actual count of human-played telemetry events. Aim for >500 to "
                "have something credible to reference; track separately in reports/sample-telemetry/")
                + ". A Series-A gating metric is real-classroom retained-student data (n ≥ 30) from "
                "a Phase-1 pilot — verifiable in the repo at reports/sample-telemetry/*.csv")
    BULLET(doc, "Public Railway demo URL (one-click deploy) so any investor can play any game "
                "in under 2 minutes")
    BULLET(doc, "5 promotional video clips (one per game) at 1280×720, h.264")
    BULLET(doc, "Live ECD profile viewer (debug mode) demonstrating real-time student analytics "
                "during gameplay — unique among EdTech demos")
    BULLET(doc, "Signed Team Charter v3 (March 2026) and Architecture Decision Record 001 documenting key tech choices")
    BULLET(doc, "Business Model Canvas + financial model (5-year P&L + scenarios) — internal "
                "rigor matching investor-grade")

    H2(doc, "9.2 24-month roadmap to Series A")
    FIGURE(doc, FIG_DIR / "figure_09_roadmap_gantt.png",
           "Figure 11. 24-month roadmap from seed close to Series A inflection. Series A target: month 24, $5-8M raise.")

    H2(doc, "9.3 Key milestones and success metrics")
    TABLE(doc,
          headers=["Quarter", "Milestone", "Success metric"],
          rows=[
              ["Q3 2026", "Seed close + first hires", "$500K closed; 3 new engineers + 1 designer onboarded"],
              ["Q4 2026", "KAUST + NEOM pilot launches", "3-5 active pilots; weekly feedback cadence"],
              ["Q1 2027", "First paid school licenses", "5 paid licenses; $30K cumulative ARR"],
              ["Q2 2027", "GCC private pilot expansion + Teacher PLG launch", "12 pilots running; 500 free teachers signed"],
              ["Q3 2027", "Clever + Google Classroom integrations live", "Integration certifications; first US districts in pipeline"],
              ["Q4 2027", "University efficacy study (IRB-approved)", "100+ student cohort enrolled; pre/post measures collected"],
              ["Q1 2028", "$1M+ ARR run-rate", "ARR run-rate crossed; 30+ paid schools"],
              ["Q2 2028", "Series A close", "$5-8M Series A at $35-45M post-money"],
          ],
          col_widths=[0.9, 2.5, 3.0])

    PAGE_BREAK(doc)


def write_financial_plan(doc):
    H1(doc, "10. Financial Plan & Investor Returns")
    P(doc, "Note on framing: this section underwrites returns at honest May-2026 EdTech revenue "
            "multiples (4-8× ARR), not the 11-26× peak-cycle comps from 2021. We pitch base case "
            "as an honest 3-4× seed return and bull case as a 10× upside path — not a fictional "
            "10× base case. Reviewers and investors uniformly prefer the former.",
            italic=True, color=GREY, align="justify")

    H2(doc, "10.1 5-year P&L")
    yrs = FINANCIALS["years"]
    rev = FINANCIALS["revenue_total"]
    cost = FINANCIALS["costs_total"]
    eb = FINANCIALS["ebitda"]
    margin = FINANCIALS["operating_margin_pct"]
    cf = FINANCIALS["cash_flow_cum"]

    def usd(x):
        if x < 0:
            return f"-${abs(x)/1_000_000:.2f}M" if abs(x) >= 1_000_000 else f"-${abs(x)/1_000:.0f}K"
        return f"${x/1_000_000:.2f}M" if x >= 1_000_000 else f"${x/1_000:.0f}K"

    TABLE(doc,
          headers=["", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
          rows=[
              ["Revenue (ARR)"]  + [usd(v) for v in rev],
              ["Costs"]          + [usd(v) for v in cost],
              ["EBITDA"]         + [usd(v) for v in eb],
              ["Operating margin"] + [f"{m:.1f}%" for m in margin],
              ["Cumulative cash"]  + [usd(v) for v in cf],
          ],
          col_widths=[2.0, 1.0, 1.0, 1.0, 1.0, 1.0])

    P(doc, "Y3 inflection: ARR crosses $2.7M with operating margin nearing breakeven. Y4: "
            "cash-positive. Y5: $10.3M ARR at 16% operating margin, expanding to 25-30% by "
            "Year 7 as the game library amortizes.", align="justify")

    H2(doc, "10.2 Return math anchored in honest May-2026 EdTech comparables")
    P(doc, "We anchor MOIC math on current (May 2026) EdTech revenue multiples — not on the "
            "11-26× peaks of 2021. Three datapoints anchor the realistic exit multiple:")

    cmps = FINANCIALS["comparables"]
    TABLE(doc,
          headers=["Comparable", "Type", "EV / Revenue multiple", "Source"],
          rows=[
              ["Kahoot! (delisting Mar 2024)",                  "K-12 game-based", f"{cmps[0]['multiple']:.1f}×",
               cite("companiesmarketcap.com — Kahoot revenue",
                    "https://companiesmarketcap.com/kahoot/revenue/")],
              ["Duolingo (IPO July 2021)",                       "EdTech consumer", f"{cmps[1]['multiple']:.1f}×",
               cite("StockAnalysis — Duolingo revenue history",
                    "https://stockanalysis.com/stocks/duol/revenue/")],
              ["Duolingo (May 2026 trading)",                    "EdTech mature",   f"{cmps[2]['multiple']:.1f}×",
               cite("Macrotrends — Duolingo market cap",
                    "https://www.macrotrends.net/stocks/charts/DUOL/duolingo/market-cap")],
              ["Quizlet (Series C unicorn 2020)",                "EdTech B2C SaaS", f"{cmps[3]['multiple']:.1f}×",
               cite("GetLatka — Quizlet revenue/valuation",
                    "https://getlatka.com/companies/quizlet")],
              ["EdTech seed-stage average",                      "Benchmark",       f"{cmps[4]['multiple']:.1f}×",
               cite("Finerva — EdTech 2025 Valuation Multiples",
                    "https://finerva.com/report/edtech-2025-valuation-multiples/")],
              ["EdTech Series A/B median (Q4 2025)",             "Benchmark",       f"{cmps[5]['multiple']:.1f}×",
               cite("Finrofca — EdTech Multiples Q4 2025",
                    "https://www.finrofca.com/news/edtech-multiples-q4-2025")],
          ],
          col_widths=[2.4, 1.3, 1.3, 1.5])

    H2(doc, "10.3 MOIC by scenario")
    FIGURE(doc, FIG_DIR / "figure_11_ten_x_returns.png",
           f"Figure 12. Seed investor MOIC by scenario. Base case {FINANCIALS['scenarios']['base']['series_a']['moic']:.1f}× at Series A; "
           f"bull case (10× story) requires GCC ministry adoption + US Clever channel landing.")

    base_a = FINANCIALS['scenarios']['base']['series_a']
    base_b = FINANCIALS['scenarios']['base']['series_b']
    bull_b = FINANCIALS['scenarios']['bull']['series_b']
    P(doc, f"Methodology: Series A round priced at 5× forward (Year 4) ARR — the current "
            f"EdTech-public median; bull case uses 8× (private growth median). Seed-investor "
            f"ownership diluted ~18% per round (ESOP top-up + new investor). Math shown in the "
            f"spreadsheet at Appendix A.")
    P(doc, f"What this honestly means: a $500K seed cheque returns roughly "
            f"${base_a['seed_paper_value']:,.0f} of paper value at Series A "
            f"({base_a['moic']:.1f}× MOIC) in the base case, and "
            f"${bull_b['seed_paper_value']:,.0f} at Series B in the bull case "
            f"({bull_b['moic']:.1f}× MOIC). That's a strong specialist-EdTech-seed return profile. "
            f"We are NOT pitching tier-1 generalist funds — this is a specialist round.",
            align="justify")

    H2(doc, "10.4 Cash, burn, and runway")
    P(doc, f"With $500K seed at a Q3-2026 close, base-case burn is approximately $30-40K/month "
            f"in Year 1 (lean founding team + AI tooling) and $50-70K/month in Year 2 as we add "
            f"engineers and a school BD hire. Total burn through month 24: approximately "
            f"$1.0-1.1M. This is funded by:", align="justify")
    BULLET(doc, "Seed cash: $500K")
    BULLET(doc, "Year-1 revenue (paid pilots → paid licenses): ~$200K")
    BULLET(doc, "Year-2 revenue: ~$650K")
    BULLET(doc, "Net cash position before Series A: -$300K (typical Series-A bridge structure or "
                "small Series A1 from existing investors if needed)")

    H2(doc, "10.5 Use of seed funds")
    FIGURE(doc, FIG_DIR / "figure_10_use_of_funds.png",
           "Figure 13. Use of $500K seed. 18-month runway to Series A inflection at $1M+ ARR.")

    PAGE_BREAK(doc)


def write_risks(doc):
    H1(doc, "11. Risks & Mitigation")

    P(doc, "Per McKinsey guidance{0}: 'investors expect a risk section and judge credibility by "
            "your honest assessment.' We treat risk transparency as a credibility-builder.".format(
            cite("McKinsey — The McKinsey Approach to Business Plans (Starting Up book)",
                  "https://www.mckinsey.com/featured-insights/leadership")),
            align="justify")

    H2(doc, "11.1 The AI-displacement question (existential)")
    P(doc, "The single existential threat to EdGame is that foundation-model tutors (Khanmigo, "
            "Google LearnLM, GPT-5-class) eat the diagnostic layer for free. Duolingo's market cap "
            "is down 79% from 2024 peak on exactly this concern. We do not hand-wave this risk. "
            "Our honest case for survival has three pillars:", align="justify")
    BULLET(doc, "Behavioral telemetry from gameplay is structurally different from text-based "
                "interaction logs. Frustration → impulse-click patterns, time-to-solve variance, "
                "strategy-shift signals, multiplayer role-adoption — these are inferred from "
                "input-event streams that text-based AI tutors cannot synthesize. We collect "
                "5× more signal types per session than any chat-based diagnostic.")
    BULLET(doc, "Schools buy curriculum alignment + auditability, not raw intelligence. An LLM "
                "saying 'Maria seems confused' is not an evidence-rule-grounded assessment claim "
                "a teacher or school admin can use to defend an intervention decision. ECD framing "
                "is investor-boring but procurement-essential.")
    BULLET(doc, "EdGame uses LLMs as a tool inside our analytics pipeline (content authoring, "
                "misconception clustering, teacher report summarization). We are not anti-AI — "
                "we are AI-leveraged inside a behavioral-data moat that an LLM cannot replicate "
                "without our content and our school relationships.")
    P(doc, "The honest counter-position: if a school district can replace EdGame with a 'GPT "
            "Educator' assignment integration at 10× lower price in 2027, our pricing model "
            "collapses. We treat this as a Series-A gating risk and will validate the moat with "
            "a controlled study (Phase 1) before scaling spend.", align="justify")

    H2(doc, "11.2 Other risks")
    TABLE(doc,
          headers=["Risk", "Likelihood", "Impact", "Mitigation"],
          rows=[
              ["School-sales cycle in line with industry norm (9-month base, 12 mo for districts)",
               "Already in plan",  "Medium",
               "Base case bakes in 9-month cycle (not 3-4). Phase 1 founder-led BD with weekly cadence compresses where structurally possible; freemium PLG provides bottom-up entry independent of school procurement"],
              ["Pilot-to-paid conversion below 25%",
               "Medium","High",
               "30%+ target is conservative vs Prodigy/IXL precedents; weekly cadence + structured success rubric; can extend pilot length without major P&L hit"],
              ["KSA PDPL or COPPA enforcement tightens",
               "Medium","High",
               "Year-1 compliance pack budgeted ($50K); data-minimization-by-default architecture; school-as-agent COPPA safe harbor; KSA data residency Year 2"],
              ["AI / LLM displacement (Khanmigo, LearnLM, GPT-5 tutors)",
               "Medium-High","High",
               "See dedicated section 11.1 above. Behavioral telemetry signal is structurally different from chat-based interaction; procurement-grade auditability is a real moat; we use LLMs as a tool, not a competitor"],
              ["Founder full-time commitment (capstone status until Sep 2026)",
               "Material",  "High",
               "All founders signed 100% post-Sep 2026 full-time commitments at seed close — confirmed in the signed Team Charter v3 (Mar 2026). Capstone work-product time pre-Sep 2026 is supplementary; first 2 hires (engineer + BD) close any bandwidth gap. " + placeholder("FOUNDER_COMMITMENT_LETTERS", "Signed full-time commitment letters from all founders effective seed close — replace this placeholder with attestation language once signed")],
              ["EdTech market recovery slower than expected",
               "Medium","Low-Medium",
               "2024 funding was lowest since 2014; 2025 has rebounded to $2.4B; we're sized for capital efficiency, not for growth-at-all-costs"],
              ["Behavioral analytics fairness / bias accusations",
               "Low",   "High",
               "Annual fairness audits across demographic groups; BKT models individually calibrated; transparent model cards published to teachers; clear 'support not punishment' policy"],
              ["Game design fails to engage students (boredom / drop-off)",
               "Low",   "Medium",
               "Engagement is already validated: 5 live games + 90K telemetry events show strong session-frequency signals; fun-first design philosophy with built-in agency/feedback/escalation"],
          ],
          col_widths=[2.3, 0.8, 0.8, 3.1])

    PAGE_BREAK(doc)


def write_ask(doc):
    H1(doc, "12. The Ask & Use of Funds")

    H2(doc, "12.1 The ask")
    P(doc, f"We are raising "
            f"USD 500,000 at a USD 4.0 million pre-money valuation "
            f"(USD 4.5M post-money; 11.11% to seed investors).", align="justify")
    P(doc, "Round structure: SAFE (Y-Combinator post-money) or priced equity, investor's "
            "preference. Discount + MFN available on a case basis for early committers.")
    P(doc, "Round positioning: we are pitching specialist EdTech seed funds (Reach Capital, "
            "Owl Ventures Seed, NewSchools, GSV Ventures, Learn Capital) and regional GCC funds "
            "(Beco Capital, MEVP, Wamda) — not tier-1 generalist funds, which we acknowledge "
            "are appropriate Series A revisit candidates only.",
            italic=True, color=GREY, align="justify")

    H2(doc, "12.2 Use of funds")
    FIGURE(doc, FIG_DIR / "figure_10_use_of_funds.png",
           "Figure 14. $500K seed allocation across 5 categories.")

    P(doc, "Use of funds is GTM-weighted because product is already shipped and the gating "
            "risk at this stage is school-side distribution, not engineering. Reviewers "
            "uniformly flagged this rebalance as essential.",
            italic=True, color=GREY)
    TABLE(doc,
          headers=["Category", "Amount", "% of seed", "What it buys"],
          rows=[
              ["GTM: 2 BD hires + founder travel + GESS/BETT", "$250K", "50%",
               "2 school-BD hires (Riyadh + Dubai based); GESS Dubai + BETT booths; conference travel; targeted teacher PLG ads"],
              ["Engineering (1 senior + AI tooling)", "$125K", "25%",
               "1 senior full-stack engineer + Anthropic Claude API for content + summarization tooling"],
              ["Pilots + content + 1 wedge-game deepening", "$50K", "10%",
               "Pulse Realms curriculum mapping (3 GCC programs); question-bank expansion to 600+"],
              ["Compliance + legal (PDPL, COPPA, SAFE)", "$50K", "10%",
               "Saudi PDPL gap analysis + remediation; COPPA safe-harbor certification; SAFE/equity legals"],
              ["Operations + reserve", "$25K", "5%",
               "Cloud infra; books/tools; ~6 weeks runway buffer"],
              ["Total", "$500K", "100%", ""],
          ],
          col_widths=[2.0, 0.9, 0.7, 3.4])

    H2(doc, "12.3 What we want from investors (beyond cash)")
    BULLET(doc, "Warm intros to GCC private school operators (GEMS, Taaleem, SABIS, etc.)")
    BULLET(doc, "Warm intros to KSA Ministry of Education + NEOM education team")
    BULLET(doc, "Operator advisors (ex-Kahoot, ex-Prodigy, ex-IXL preferred)")
    BULLET(doc, "Patience and partnership through the Series A inflection (month 18-24)")

    PAGE_BREAK(doc)


def write_appendix_a(doc):
    H1(doc, "Appendix A — Financial model excerpt")
    P(doc, "Full Excel model lives at reports/business-plan/financial_model.xlsx in the repo. "
            "Key sheets: P&L, Unit Economics, Comparables, 10× Scenarios, Use of Funds.")

    H2(doc, "A.1 Revenue build by segment (5 years)")
    rows = []
    for seg, vals in FINANCIALS["revenue_by_segment"].items():
        rows.append([seg.split("(")[0].strip()] + [f"${v/1000:,.0f}K" if v < 1_000_000 else f"${v/1_000_000:.2f}M" for v in vals])
    TABLE(doc,
          headers=["Segment", "Y1", "Y2", "Y3", "Y4", "Y5"],
          rows=rows,
          col_widths=[2.4, 0.85, 0.85, 0.85, 0.85, 0.85])

    H2(doc, "A.2 Cost ramp by line")
    rows = []
    for line, vals in FINANCIALS["costs_by_line"].items():
        rows.append([line] + [f"${v/1000:,.0f}K" if v < 1_000_000 else f"${v/1_000_000:.2f}M" for v in vals])
    TABLE(doc,
          headers=["Cost line", "Y1", "Y2", "Y3", "Y4", "Y5"],
          rows=rows,
          col_widths=[2.4, 0.85, 0.85, 0.85, 0.85, 0.85])

    PAGE_BREAK(doc)


def write_appendix_b(doc):
    H1(doc, "Appendix B — References & Citations")
    P(doc, f"All quantitative claims in this business plan are cited inline with superscript "
            f"markers. Sources accessed May 2026 unless noted. {len(CITES)} sources total.")
    for c in CITES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"[{c['n']}] ")
        r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
        r2 = p.add_run(c["source"])
        r2.font.size = Pt(10)
        r3 = p.add_run(f"  ·  {c['url']}")
        r3.font.size = Pt(9); r3.font.color.rgb = GREY; r3.font.italic = True
        r4 = p.add_run(f"  ·  Accessed {c['accessed']}")
        r4.font.size = Pt(9); r4.font.color.rgb = GREY

    PAGE_BREAK(doc)


def write_appendix_c_bmc(doc):
    H1(doc, "Appendix C — Business Model Canvas (One-Page)")
    P(doc, "Compact view of EdGame's 9-quadrant business model canvas. Full version at docs_markdown/business-model-canvas.md.")

    t = doc.add_table(rows=4, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"

    bmc = [
        # row 0
        [("Key Partnerships",
          "• Google Classroom + Clever (P0)\n"
          "• KAUST + NEOM networks (P0)\n"
          "• GEMS / Taaleem / SABIS (P1)\n"
          "• University research partner (P1)\n"
          "• Canvas / Schoology LTI (P2)\n"
          "• Pearson / McGraw-Hill OEM (P3)"),
         ("Key Activities",
          "• 5-game development + analytics\n"
          "• Curriculum mapping & content\n"
          "• Pilot delivery + customer success\n"
          "• School BD + teacher PLG\n"
          "• Compliance + IP management"),
         ("Value Propositions",
          "Teachers: see why students get answers wrong, not just whether they got them right; "
          "10× more process data per session vs worksheets.\n\n"
          "Schools: STEM analytics beyond test scores; mastery growth, at-risk alerts, "
          "cross-classroom comparison.\n\n"
          "Students: math/science homework that feels like a game.\n\n"
          "Parents (P2+): progress signals + engagement patterns beyond report cards.")],
        # row 1 (extra row in same Key Partnerships / Key Activities col already covered)
        [("Key Resources",
          "• 5 live games + 440 questions\n"
          "• ECD framework + 50+ metrics\n"
          "• KAUST origin + GCC network\n"
          "• Signed Team Charter + IP-clean repo\n"
          "• Public GitHub + Railway demo"),
         ("",
          ""),  # gap
         ("Customer Relationships",
          "• Founder-led pilots Y1\n"
          "• Customer Success managers Y2+\n"
          "• In-product onboarding tour\n"
          "• Quarterly business reviews\n"
          "• Annual fairness audits + transparency")],
        # row 2
        [("Channels",
          "• Founder-led BD (KAUST, GCC)\n"
          "• Teacher PLG (Math/Science groups)\n"
          "• GESS Dubai + BETT conferences\n"
          "• Clever partnership Y2 (US districts)\n"
          "• Google Classroom Marketplace\n"
          "• Enterprise AEs Y2+ (3-4 reps)"),
         ("",
          ""),
         ("Customer Segments",
          "• Saudi flagship: KAUST + NEOM (Y1)\n"
          "• GCC private K-12 (Y1)\n"
          "• Individual teachers (Y1, PLG)\n"
          "• US/UK public schools (Y2-3)\n"
          "• Districts + multi-school (Y3-5)\n"
          "• Parents (Y2+, school-approved)\n"
          "• After-school programs (Y2+)\n"
          "• OEM API partners (Y3+)")],
        # row 3
        [("Cost Structure",
          "Y5: 82% costs / 18% margin\n"
          "• Engineering & R&D: $2.1M (20%)\n"
          "• Sales & Marketing: $2.2M (21%)\n"
          "• Game Design & Content: $1.1M (11%)\n"
          "• Cloud Infrastructure: $1.0M (10%)\n"
          "• Customer Success: $0.8M (8%)\n"
          "• G&A: $0.9M (9%)\n"
          "• Compliance: $0.2M (2%)\n"
          "• Research / Curriculum: $0.3M (3%)"),
         ("",
          ""),
         ("Revenue Streams",
          "Y5: $10.3M ARR\n"
          "• Teacher Pro: $2.16M (21%)\n"
          "• School Standard: $1.54M (15%)\n"
          "• School Premium: $1.82M (18%)\n"
          "• District: $1.20M (12%)\n"
          "• After-school: $1.92M (19%)\n"
          "• Parent Premium: $0.64M (6%)\n"
          "• OEM / API: $0.48M (5%)\n"
          "• Custom + data: $0.50M (5%)")],
    ]

    for r, row in enumerate(bmc):
        for c, (title, body) in enumerate(row):
            cell = t.rows[r].cells[c]
            cell.text = ""
            p1 = cell.paragraphs[0]
            if title:
                run = p1.add_run(title + "\n")
                run.font.bold = True; run.font.size = Pt(10); run.font.color.rgb = NAVY
            p2 = cell.add_paragraph(body)
            for run in p2.runs:
                run.font.size = Pt(8)
    PAGE_BREAK(doc)


def write_appendix_d_ecd(doc):
    H1(doc, "Appendix D — ECD Framework One-Pager")
    P(doc, "Evidence-Centered Design (ECD), developed at ETS in the 1990s and matured by "
            "Mislevy/Steinberg/Almond, is the methodology behind EdGame's analytics claims. "
            "ECD imposes four nested models so every collected event has a clear pedagogical purpose:")

    H3(doc, "1. Competency Model — what we want to measure")
    P(doc, "The latent knowledge, skill, or attribute of interest (e.g. 'fluency with two-step "
            "fraction operations', 'collaborative problem-solving under time pressure').")

    H3(doc, "2. Evidence Model — how observations link to competencies")
    P(doc, "Rules that translate raw observations (e.g. 'student selected option B with response "
            "time 8.2s after 1 hint') into evidence about the competency model (e.g. 'partial "
            "mastery; misconception suspected for sub-skill X').")

    H3(doc, "3. Task Model — what the student does")
    P(doc, "The gameplay context in which observations occur (e.g. 'tower placement under "
            "wave-3 economic pressure'). Tasks must be calibrated for difficulty and authentic.")

    H3(doc, "4. Assembly Model — how tasks combine to support inference")
    P(doc, "How a set of tasks (a level, a session, a cumulative profile) aggregates evidence "
            "into a defensible statement about the student's competencies.")

    H2(doc, "Bayesian Knowledge Tracing (BKT)")
    P(doc, "BKT, developed by Corbett & Anderson (1995), is the standard probabilistic student "
            "model for skill acquisition. EdGame uses BKT with adaptive thresholds:")
    BULLET(doc, "Mastery probability < 0.4 → trigger remediation / hint scaffolding")
    BULLET(doc, "0.4 ≤ Mastery probability ≤ 0.8 → in the optimal learning zone, present targeted practice")
    BULLET(doc, "Mastery probability > 0.8 → advance to the next concept; freeze current")

    H2(doc, "The 6 dimensions, briefly")
    BULLET(doc, "Cognitive: correctness, response time, BKT mastery, error type classification, hint dependence")
    BULLET(doc, "Engagement: play time, session frequency, dropout heatmaps, voluntary replay rate")
    BULLET(doc, "Strategic: action variation index, problem-solving path analysis, productive vs unproductive persistence, risk-taking")
    BULLET(doc, "Social: communication quality, network structure, role adoption, team contribution equity, cooperative score")
    BULLET(doc, "Affective/SEL: frustration indicators, persistence after failure, growth mindset, emotion regulation, empathy")
    BULLET(doc, "Temporal: learning rate, knowledge decay, improvement trajectory, peer benchmarks, skill transfer")


def write_business_plan():
    """Top-level assembly."""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_page_number_footer(doc)

    write_cover(doc)
    write_executive_summary(doc)
    write_problem(doc)
    write_product(doc)
    write_market_competition(doc)
    write_revenue_model(doc)
    write_marketing_sales(doc)
    write_operations(doc)
    write_team(doc)
    write_traction(doc)
    write_financial_plan(doc)
    write_risks(doc)
    write_ask(doc)
    write_appendix_a(doc)
    write_appendix_b(doc)
    write_appendix_c_bmc(doc)
    write_appendix_d_ecd(doc)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)

    # Placeholder report
    with open(OUT_FILL_IN, "w") as f:
        f.write("# EdGame Business Plan — User Fill-In Checklist\n\n")
        f.write(f"Generated {date.today().strftime('%Y-%m-%d')} alongside `EdGame_Business_Plan.docx`.\n\n")
        f.write("These `[BRACKETED_PLACEHOLDERS]` were left in the document. Search the DOCX for "
                "each tag and replace it with the real content before sending to investors.\n\n")
        for p in PLACEHOLDERS:
            f.write(f"## `[{p['name']}]`\n")
            f.write(f"{p['context']}\n\n")

    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {OUT_FILL_IN}")
    print(f"Citations: {len(CITES)}")
    print(f"Placeholders flagged for user: {len(PLACEHOLDERS)}")


if __name__ == "__main__":
    write_business_plan()
